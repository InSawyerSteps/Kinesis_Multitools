"""Recruiting MCP Server
=======================
Standalone FastAPI application that will expose recruiting-specific endpoints
(indexing resumes, semantic search, hybrid search, JD / resume parsing).

Phase 1: Boilerplate only – real business logic will be copied in subsequent
commits. All endpoints currently return 501 (Not Implemented) so we can verify
that the server spins up correctly on port 8001.
"""
from __future__ import annotations

import logging
import os
import pathlib
import platform
import sys
from typing import List, Dict, Optional, Any

# Third-party helpers used by the recruiting tools
import requests  # for Ollama embeddings calls
import uuid
import json

try:
    from thefuzz import fuzz  # type: ignore
except ImportError:
    fuzz = None

from fastapi import FastAPI, HTTPException, status
from pydantic import BaseModel
import uvicorn

# ---------------------------------------------------------------------------
# Basic path & logging setup – mirrors original toolz.py pattern
# ---------------------------------------------------------------------------
SCRIPT_PATH = pathlib.Path(__file__).resolve()
PROJECT_ROOT = SCRIPT_PATH.parent.parent  # assume repo/{src,this_file}
os.chdir(PROJECT_ROOT)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("recruiting_mcp_server")
logger.setLevel(logging.DEBUG)
logger.info("Recruiting MCP Server starting – Python %s on %s", platform.python_version(), platform.system())

# ---------------------------------------------------------------------------
# Configuration for Resume Storage & ChromaDB
# ---------------------------------------------------------------------------
PROJECT_RESUME_PATHS: Dict[str, pathlib.Path] = {
    "RecruitingDemo": pathlib.Path(r"C:\Projects\RecruitingDemo\resumes"),
    "Rec_demo": pathlib.Path(r"C:\Projects\MCP Server\Samples"),
}

CHROMA_DATA_PATH = "chroma_db_data"
RESUME_COLLECTION_NAME = "resumes"

try:
    import chromadb  # type: ignore

    chroma_client = chromadb.PersistentClient(path=str(PROJECT_ROOT / CHROMA_DATA_PATH))
    resume_collection = chroma_client.get_or_create_collection(name=RESUME_COLLECTION_NAME)
    logger.info("ChromaDB resume collection initialised at %s", PROJECT_ROOT / CHROMA_DATA_PATH)
except Exception as chroma_e:
    resume_collection = None
    logger.error("Failed to initialise ChromaDB resume collection: %s", chroma_e, exc_info=True)

# ---------------------------------------------------------------------------
# FastAPI app instance
# ---------------------------------------------------------------------------
app = FastAPI(
    title="Recruiting MCP Server",
    description="Dedicated server for resume and job description analysis tools.",
    version="1.0.0",
)

# ---------------------------------------------------------------------------
# Pydantic request models (stubs for now)
# ---------------------------------------------------------------------------
class IndexResumeRequest(BaseModel):
    project_name: str
    file_path: str
    chunk_size: int = 500
    chunk_overlap: int = 50


class SemanticSearchRequest(BaseModel):
    query: str
    max_results: int = 5


class HybridSearchRequest(BaseModel):
    query: str
    core_technologies: List[str]
    max_results: int = 10
    tech_aliases: Optional[Dict[str, List[str]]] = None
    required_technologies: Optional[List[str]] = None
    score_weights: Optional[Dict[str, float]] = None
    fuzzy_threshold: int = 85


class ParseJDRequest(BaseModel):
    file_path: Optional[str] = None
    raw_text: Optional[str] = None


class ParseResumeRequest(BaseModel):
    project_name: str
    file_path: str


# ---------------------------------------------------------------------------
# Temporary placeholder endpoints – will be replaced with real logic in Phase 2
# ---------------------------------------------------------------------------
# ---------------------------------------------------------------------------
# Resume Indexing (migrated & simplified)
# ---------------------------------------------------------------------------

def index_resume(project_name: str, file_path: str, chunk_size: int = 500, chunk_overlap: int = 50) -> dict:  # noqa: C901
    """Index a resume into ChromaDB embedding collection."""
    logger.info("[index_resume] project=%s, file=%s", project_name, file_path)
    if resume_collection is None:
        return {"status": "error", "message": "ChromaDB resume collection not initialised."}

    # Parse resume text first
    parse_result = parse_resume(project_name, file_path)
    if parse_result.get("status") != "success":
        return {"status": "error", "message": parse_result.get("message")}

    raw_text: str = parse_result.get("raw_text", "").strip()
    meta_base: dict[str, Any] = parse_result.get("extracted_fields", {})
    if not raw_text:
        return {"status": "error", "message": "No text to index."}

    # --- Chunking (reuse section heading detection) ---
    headings = [
        r"SUMMARY", r"EXPERIENCE", r"EDUCATION", r"SKILLS", r"PROJECTS", r"CERTIFICATION",
    ]
    heading_regex = re.compile(rf"^({'|'.join(headings)})[\s:]*$", re.MULTILINE | re.IGNORECASE)
    matches = list(heading_regex.finditer(raw_text))
    chunks: List[str] = []
    if matches:
        positions = [m.start() for m in matches] + [len(raw_text)]
        for i in range(len(positions) - 1):
            chunk = raw_text[positions[i]:positions[i + 1]].strip()
            if chunk:
                chunks.append(chunk)
    else:
        start = 0
        while start < len(raw_text):
            end = min(start + chunk_size, len(raw_text))
            chunk = raw_text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            start += chunk_size - chunk_overlap

    embeddings: List[List[float]] = []
    embedding_errors: List[str] = []
    for chunk in chunks:
        try:
            resp = requests.post(
                "http://localhost:11434/api/embeddings",
                json={"model": "nomic-embed-text", "prompt": chunk},
                timeout=30,
            )
            resp.raise_for_status()
            emb = resp.json().get("embedding")
            if not emb:
                raise ValueError("Empty embedding")
            embeddings.append(emb)
        except Exception as e:
            embeddings.append([0.0])
            embedding_errors.append(str(e))

    ids = [str(uuid.uuid4()) for _ in chunks]

    def _scalar(v: Any):
        if isinstance(v, (str, int, float, bool)) or v is None:
            return v
        try:
            return json.dumps(v, ensure_ascii=False)
        except Exception:
            return str(v)

    metadatas = [
        {
            "file_path": file_path,
            "project_name": project_name,
            "chunk_index": i,
            **{k: _scalar(v) for k, v in meta_base.items()},
        }
        for i in range(len(chunks))
    ]

    try:
        resume_collection.add(documents=chunks, embeddings=embeddings, metadatas=metadatas, ids=ids)
    except Exception as e:
        return {"status": "error", "message": f"ChromaDB add failed: {e}"}

    return {
        "status": "success",
        "chunks_indexed": len(chunks),
        "embedding_errors": embedding_errors,
    }


# ---------------------------------------------------------------------------
# Semantic Search over resumes
# ---------------------------------------------------------------------------

def semantic_search_in_resumes(query: str, max_results: int = 5) -> dict:
    if resume_collection is None:
        return {"results": [], "error": "ChromaDB resume collection not initialised."}
    try:
        resp = requests.post(
            "http://localhost:11434/api/embeddings",
            json={"model": "nomic-embed-text", "prompt": query},
            timeout=30,
        )
        resp.raise_for_status()
        q_emb = resp.json().get("embedding")
        if not q_emb:
            raise ValueError("No embedding returned")
    except Exception as e:
        return {"results": [], "error": str(e)}

    search = resume_collection.query(
        query_embeddings=[q_emb],
        n_results=max_results,
        include=["documents", "metadatas", "distances"],
    )
    docs = search.get("documents", [[]])[0]
    metas = search.get("metadatas", [[]])[0]
    dists = search.get("distances", [[]])[0]

    results: List[dict[str, Any]] = []
    for i in range(len(docs)):
        sim = 1.0 / (1.0 + dists[i]) if dists[i] is not None else 0.0
        results.append({
            "file_path": metas[i].get("file_path", "Unknown"),
            "chunk_content": docs[i],
            "similarity_score": round(sim, 4),
            "metadata": metas[i],
        })
    return {"results": results}


# ---------------------------------------------------------------------------
# Hybrid Search (semantic + tech keyword/alias + optional fuzzy)
# ---------------------------------------------------------------------------

def hybrid_search_recruiting_data(
    query: str,
    core_technologies: List[str],
    max_results: int = 10,
    tech_aliases: Optional[Dict[str, List[str]]] = None,
    required_technologies: Optional[List[str]] = None,
    score_weights: Optional[Dict[str, float]] = None,
    fuzzy_threshold: int = 85,
) -> dict:  # noqa: C901
    """Rank candidates combining semantic similarity with tech keyword matching."""
    if not query or not core_technologies:
        return {"status": "error", "message": "query and core_technologies required"}
    if fuzzy_threshold <= 100 and fuzz is None:
        return {"status": "error", "message": "fuzzy matching requested but thefuzz not installed"}

    tech_aliases = tech_aliases or {}
    required_technologies = set(required_technologies) if required_technologies else set()
    score_weights = score_weights or {"semantic": 0.6, "tech": 0.4}

    sem = semantic_search_in_resumes(query, max_results=max_results * 2)
    if "results" not in sem:
        return {"status": "error", "message": sem.get("error", "semantic search failed")}

    ranked: List[Dict[str, Any]] = []
    for cand in sem["results"]:
        text = cand.get("chunk_content", "").lower()
        matched, missing, matched_alias = set(), set(core_technologies), {}
        for tech in core_technologies:
            aliases = {tech.lower(), *(a.lower() for a in tech_aliases.get(tech, []))}
            alias_found = next((a for a in aliases if a in text), None)
            if not alias_found and 0 <= fuzzy_threshold <= 100 and fuzz is not None:
                alias_found = next((a for a in aliases if fuzz.partial_ratio(a, text) >= fuzzy_threshold), None)
            if alias_found:
                matched.add(tech)
                matched_alias[tech] = alias_found
                if tech in missing:
                    missing.remove(tech)
        if required_technologies and not required_technologies.issubset(matched):
            continue
        tech_score = len(matched) / len(core_technologies)
        sem_score = cand.get("similarity_score", 0.0)
        final_score = sem_score * score_weights["semantic"] + tech_score * score_weights["tech"]
        cand.update(
            {
                "matched_techs": list(matched),
                "missing_techs": list(missing),
                "matched_aliases": matched_alias,
                "tech_match_score": round(tech_score, 4),
                "final_score": round(final_score, 4),
            }
        )
        ranked.append(cand)

    ranked.sort(key=lambda x: x["final_score"], reverse=True)
    return {"results": ranked[:max_results]}


# ---------------------------------------------------------------------------
# Job Description Parsing (simplified version of original)
# ---------------------------------------------------------------------------

def parse_job_description(file_path: str | None = None, raw_text: str | None = None) -> dict:
    if not (file_path or raw_text):
        return {"status": "error", "message": "Provide file_path or raw_text"}

    text = raw_text or ""
    if file_path:
        try:
            ext = pathlib.Path(file_path).suffix.lower()
            if ext in {".txt", ".md"}:
                text = pathlib.Path(file_path).read_text(encoding="utf-8", errors="replace")
            elif ext == ".pdf" and extract_pdf_text:
                text = extract_pdf_text(str(file_path))
            elif ext == ".docx" and docx:
                text = "\n".join(p.text for p in docx.Document(str(file_path)).paragraphs)
            else:
                return {"status": "error", "message": f"Unsupported extension {ext}"}
        except Exception as e:
            return {"status": "error", "message": f"failed to read file: {e}"}

    if not text.strip():
        return {"status": "error", "message": "empty JD text"}

    # naive extraction
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    title = lines[0] if lines else "Job Description"

    skills_found = re.findall(r"python|fastapi|java|javascript|aws|docker|kubernetes", text, flags=re.I)
    skills_unique = sorted(set([s.lower() for s in skills_found]))

    years_match = re.search(r"(\d+)\+?\s+years", text, flags=re.I)
    min_exp = int(years_match.group(1)) if years_match else None

    return {
        "status": "success",
        "title": title,
        "required_skills": skills_unique,
        "min_experience": min_exp,
        "raw_text": text,
    }


# ---------------------------------------------------------------------------
# API endpoints for the above functions
# ---------------------------------------------------------------------------
@app.post("/index-resume")
async def api_index_resume(req: IndexResumeRequest):
    result = index_resume(req.project_name, req.file_path, req.chunk_size, req.chunk_overlap)
    if result.get("status") != "success":
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=result.get("message"))
    return result


@app.post("/semantic-search-resumes")
async def api_semantic_search(req: SemanticSearchRequest):
    result = semantic_search_in_resumes(req.query, req.max_results)
    if result.get("error"):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=result.get("error"))
    return result








@app.post("/parse-jd")
async def api_parse_jd(req: ParseJDRequest):
    result = parse_job_description(req.file_path, req.raw_text)
    if result.get("status") != "success":
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=result.get("message"))
    return result


# ---------------------------------------------------------------------------
# Hybrid Search Endpoint
# ---------------------------------------------------------------------------
@app.post("/hybrid-search")
async def api_hybrid_search(req: HybridSearchRequest):
    result = hybrid_search_recruiting_data(
        query=req.query,
        core_technologies=req.core_technologies,
        max_results=req.max_results,
        tech_aliases=req.tech_aliases,
        required_technologies=req.required_technologies,
        score_weights=req.score_weights,
        fuzzy_threshold=req.fuzzy_threshold,
    )
    if result.get("status") == "error":
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=result.get("message"))
    return result


# ---------------------------------------------------------------------------
# Resume Parsing Logic (migrated from toolz.py)
# ---------------------------------------------------------------------------
import mimetypes
import chardet

try:
    import docx  # type: ignore
except ImportError:
    docx = None

try:
    from pdfminer.high_level import extract_text as extract_pdf_text  # type: ignore
except ImportError:
    extract_pdf_text = None

try:
    import PyPDF2  # type: ignore
except ImportError:
    PyPDF2 = None


def parse_resume(project_name: str, file_path: str) -> dict:  # noqa: C901 – long, copied verbatim with minimal tweaks
    """Parses a resume file and extracts raw text plus key fields.

    Args:
        project_name: Recruiting project name present in PROJECT_RESUME_PATHS.
        file_path: Absolute path to the resume file.

    Returns:
        Dict with status, raw_text, extracted_fields (emails, phone_numbers, etc.), and message.
    """
    logger.info("Executing parse_resume for project=%s file=%s", project_name, file_path)

    if project_name not in PROJECT_RESUME_PATHS:
        return {
            "status": "error",
            "file_path": file_path,
            "raw_text": "",
            "extracted_fields": {},
            "message": f"Invalid project name: {project_name}",
        }

    resume_root = PROJECT_RESUME_PATHS[project_name]
    try:
        abs_resume_root = resume_root.resolve()
        abs_file_path = pathlib.Path(file_path).resolve()
        if not abs_file_path.is_file() or not str(abs_file_path).startswith(str(abs_resume_root)):
            return {
                "status": "error",
                "file_path": file_path,
                "raw_text": "",
                "extracted_fields": {},
                "message": "File not found or not within allowed path.",
            }
    except Exception as e:
        return {
            "status": "error",
            "file_path": file_path,
            "raw_text": "",
            "extracted_fields": {},
            "message": f"Path validation error: {e}",
        }

    ext = pathlib.Path(file_path).suffix.lower()
    mimetype, _ = mimetypes.guess_type(str(file_path))
    raw_text: str = ""
    error: Optional[str] = None

    try:
        if ext == ".pdf" or (mimetype and "pdf" in mimetype):
            if extract_pdf_text:
                try:
                    raw_text = extract_pdf_text(str(file_path))
                except Exception as pdfminer_e:
                    if PyPDF2:
                        try:
                            with open(file_path, "rb") as f:
                                reader = PyPDF2.PdfReader(f)
                                raw_text = "\n".join(page.extract_text() or "" for page in reader.pages)
                        except Exception as pypdf_e:
                            error = f"PDF extraction failed: {pypdf_e}"
                    else:
                        error = f"pdfminer failed: {pdfminer_e}"
            elif PyPDF2:
                try:
                    with open(file_path, "rb") as f:
                        reader = PyPDF2.PdfReader(f)
                        raw_text = "\n".join(page.extract_text() or "" for page in reader.pages)
                except Exception as pypdf_e:
                    error = f"PDF extraction failed: {pypdf_e}"
            else:
                error = "No PDF extraction library available. Install pdfminer.six or PyPDF2."
        elif ext == ".docx" or (mimetype and "word" in mimetype):
            if docx:
                try:
                    doc = docx.Document(str(file_path))
                    raw_text = "\n".join(p.text for p in doc.paragraphs)
                except Exception as docx_e:
                    error = f"DOCX extraction failed: {docx_e}"
            else:
                error = "python-docx not installed."
        elif ext in {".txt", ".md"} or (mimetype and ("text" in mimetype or "markdown" in mimetype)):
            try:
                with open(file_path, "rb") as f:
                    raw_bytes = f.read()
                detected = chardet.detect(raw_bytes)
                encoding = detected.get("encoding") or "utf-8"
                raw_text = raw_bytes.decode(encoding, errors="replace")
            except Exception as txt_e:
                error = f"Text extraction failed: {txt_e}"
        else:
            return {
                "status": "unsupported",
                "file_path": file_path,
                "raw_text": "",
                "extracted_fields": {},
                "message": f"Unsupported file type: {ext}",
            }
    except Exception as e:  # pragma: no cover
        return {
            "status": "error",
            "file_path": file_path,
            "raw_text": "",
            "extracted_fields": {},
            "message": f"Extraction error: {e}",
        }

    if error:
        return {
            "status": "error",
            "file_path": file_path,
            "raw_text": raw_text,
            "extracted_fields": {},
            "message": error,
        }

    if not raw_text.strip():
        return {
            "status": "error",
            "file_path": file_path,
            "raw_text": "",
            "extracted_fields": {},
            "message": "No text extracted from file.",
        }

    # Basic extraction – placeholder for more advanced NLP logic
    emails = re.findall(r"[\w\.-]+@[\w\.-]+", raw_text)
    phones = re.findall(r"\+?\d[\d\s\-]{7,}\d", raw_text)

    extracted_fields = {
        "emails": emails,
        "phone_numbers": phones,
        "potential_skills": [],
        "names": [],
        "summary": " ".join(raw_text.split()[:40]) + ("..." if len(raw_text.split()) > 40 else ""),
        "pyresparser": {},
    }

    # Try pyresparser if available
    if ResumeParser is not None:
        try:
            data = ResumeParser(str(file_path)).get_extracted_data()
            extracted_fields["pyresparser"] = data or {}
            if data and data.get("skills"):
                extracted_fields["potential_skills"] = data["skills"]
            if data and data.get("name"):
                extracted_fields["names"] = [data["name"]]
        except Exception as rp_e:
            logger.warning("pyresparser failed: %s", rp_e)

    return {
        "status": "success",
        "file_path": file_path,
        "raw_text": raw_text,
        "extracted_fields": extracted_fields,
        "message": "Parsed successfully.",
    }


# ---------------------------------------------------------------------------
# API endpoint – uses the above function
# ---------------------------------------------------------------------------
@app.post("/parse-resume")
async def api_parse_resume(request: ParseResumeRequest):
    result = parse_resume(request.project_name, request.file_path)
    if result.get("status") != "success":
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=result.get("message"))
    return result


# ---------------------------------------------------------------------------
# Entry-point – run with: python recruiting_mcp_server.py
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    uvicorn.run("recruiting_mcp_server:app", host="127.0.0.1", port=8001, reload=False)
