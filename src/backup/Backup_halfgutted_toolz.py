import os
import pathlib
import sys
import platform
import logging
import traceback # For detailed error logging
import re # Added for new regex-based tools

# --- Define project_root safely as the parent of src (like main.py) ---
project_root = None # Initialize to None
try:
    # Get the directory containing this script (src)
    script_path = pathlib.Path(__file__).resolve()
    script_dir = script_path.parent
    # Assume the project root is one level above the 'src' directory
    project_root = script_dir.parent
    # DEBUG: Print the calculated project_root immediately
    print(f"MCP SERVER DEBUG: Calculated project_root as: {project_root}", file=sys.stderr)
except Exception as path_e:
    # Handle error defining paths separately
    print(f"MCP SERVER CRITICAL ERROR: Failed to calculate project_root based on __file__ '{__file__ if '__file__' in locals() else 'undefined'}'. Error: {path_e}", file=sys.stderr)
    sys.exit(1) # Exit if we can't even determine the project root reliably

# --- Explicitly set CWD using the calculated project_root ---
# Proceed only if project_root was successfully defined (should always be true unless exit above)
if project_root and project_root.is_dir(): # Extra check if it's a directory
    try:
        os.chdir(project_root)
        # Use print for immediate feedback before logging is configured
        print(f"MCP SERVER: Successfully changed CWD to: {os.getcwd()}", file=sys.stderr)
    except Exception as chdir_e:
        # Print error if changing CWD fails
        print(f"MCP SERVER CRITICAL ERROR: Failed to change CWD to '{project_root}'. Error: {chdir_e}", file=sys.stderr)
        sys.exit(1) # Exit if CWD change fails, as relative paths will be wrong
elif project_root:
    print(f"MCP SERVER CRITICAL ERROR: Calculated project_root '{project_root}' is not a valid directory.", file=sys.stderr)
    sys.exit(1)
else:
    # This case should not be reached if the first try/except worked or exited
    print(f"MCP SERVER CRITICAL ERROR: Cannot change CWD because project_root is not set.", file=sys.stderr)
    sys.exit(1)
# --- End CWD setting ---

# --- Now continue with the original imports and script logic ---
# Ensure we use the installed mcp package, not local file
import sys
sys.path.insert(0, 'C:\Projects\MCP Server\.venv\Lib\site-packages')
from fastmcp import FastMCP
import ollama
from markdown_it import MarkdownIt
from markdown_it.token import Token # For type hinting
from typing import List, Dict, Optional
try:
    from thefuzz import fuzz
except ImportError:
    fuzz = None  # Will check at runtime

# Configure basic logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger("mcp_server_module") # Give logger a specific name
logger.setLevel(logging.DEBUG)

# --- Tree-sitter Setup ---
# Note: For simplicity in this version, we'll skip the tree-sitter integration
# and make the tool gracefully report that syntax analysis is not available
# A full implementation would require more setup and potentially building language libs
DART_LANGUAGE = None
dart_parser = None
logger.warning("Tree-sitter Dart language support is not currently configured. The get_dart_definition tool will operate in limited mode.")
# --- End Tree-sitter Setup ---

# Create the FastMCP server instance BEFORE any @mcp.tool() decorators
logger.info("DEBUG: Preparing to instantiate FastMCP...") # <-- ADD THIS
mcp = FastMCP(
    name="Local Project Context Server",
    title="MCP Server (MVP + Doc Tools)",
    description="Provides context from local project documentation (incl. flows, wireframes, styles) and potentially databases.",
    version="0.2.0"
)
logger.info(f"DEBUG: FastMCP instance created. Type: {type(mcp)}, Has version attr: {hasattr(mcp, 'version')}") # <-- ADD THIS
if hasattr(mcp, 'version'): # <-- ADD THIS BLOCK
    logger.info(f"DEBUG: mcp.version value is: {mcp.version}")
else:
    logger.warning("DEBUG: mcp object DOES NOT have 'version' attribute immediately after creation!")

# --- Configuration for Project Documentation Paths ---
PROJECT_DOC_PATHS = {
    "ParentBuddy": pathlib.Path(r"C:\Projects\ParentBuddy\Docs"),
    "DreamApp": pathlib.Path(r"C:\Projects\DreamApp\Docs"),
    "MCPServer": pathlib.Path(r"C:\Projects\MCP Server\Docs"),
    "RecruitingRAG": pathlib.Path(r"C:\Projects\MCP Server\Docs")
}
# --- End Configuration ---

# --- Configuration for Project Source File Paths (Dart) ---
# You can add more projects or adjust these as needed.
PROJECT_SRC_PATHS = {
    "ParentBuddy": pathlib.Path(r"C:\Projects\ParentBuddy\lib"),
    "DreamApp": pathlib.Path(r"C:\Projects\DreamApp\lib")
}
# --- End Configuration ---

# --- Configuration for Project Drift File Locations ---
# Points to the directory from where to start searching for .drift files
PROJECT_DRIFT_PATHS = {
    "ParentBuddy": pathlib.Path(r"C:\Projects\ParentBuddy"), # Search from project root
    "DreamApp": pathlib.Path(r"C:\Projects\DreamApp")      # Search from project root
    # Add other projects if they use Drift
}
# --- End Configuration ---

# --- Configuration for Resume File Locations ---
# Points to the directory from where to start searching for resume files
# Recruiting resume paths moved to recruiting_mcp_server
# PROJECT_RESUME_PATHS = {
#     "RecruitingDemo": pathlib.Path(r"C:\Projects\RecruitingDemo\resumes"),
#     "Rec_demo": pathlib.Path(r"C:\Projects\MCP Server\Samples"),
#     # Add other projects or folders as needed
}
# --- End Configuration ---


# --- ChromaDB Setup ---
CHROMA_DATA_PATH = "chroma_db_data" # Relative path within project root
COLLECTION_NAME = "project_documentation"
# Recruiting collection name moved to recruiting_mcp_server
# RESUME_COLLECTION_NAME = "resumes"

# Initialize ChromaDB client and collections globally
chroma_client = None
collection = None
resume_collection = None
try:
    import chromadb
    # Ensure the path is resolved relative to the now-correct CWD
    abs_chroma_path = project_root / CHROMA_DATA_PATH
    chroma_client = chromadb.PersistentClient(path=str(abs_chroma_path))
    logger.info(f"ChromaDB client initialized. Data path: {abs_chroma_path}")
    collection = chroma_client.get_or_create_collection(name=COLLECTION_NAME)
    logger.info(f"ChromaDB collection '{COLLECTION_NAME}' loaded/created.")
    # Resume collection handled by recruiting_mcp_server
    resume_collection = None  # disabled here
    # logger.info("Resume collection initialised in original server – now disabled here")
except Exception as e:
    logger.error(f"Failed to initialize ChromaDB: {e}", exc_info=True)
# --- End ChromaDB Setup ---

# --- Markdown Parser Setup ---
md_parser = MarkdownIt()
# --- End Parser Setup ---


# --- Resume Indexing Tool ---
import uuid

# Recruiting tool moved to recruiting_mcp_server – decorator removed
def index_resume(project_name: str, file_path: str, chunk_size: int = 500, chunk_overlap: int = 50) -> dict:
    """
    Indexes a resume: parses, chunks (by section/heading if possible), embeds, and stores in ChromaDB ('resumes' collection) with metadata.
    Args:
        project_name: The name of the recruiting project (must be in PROJECT_RESUME_PATHS).
        file_path: The absolute path to the resume file.
        chunk_size: Max characters per chunk (default 500, used if fallback to fixed-size chunking).
        chunk_overlap: Overlap between chunks (default 50, used if fallback to fixed-size chunking).
    Returns:
        Dict with status, message, number of chunks indexed, and errors if any.
    Notes:
        - Tries to split resume by common section headings (e.g., SUMMARY, EXPERIENCE, EDUCATION, SKILLS, PROJECTS).
        - If no headings are found, falls back to fixed-size overlapping chunking.
    """
    logger.info(f"Starting index_resume for project: {project_name}, file: {file_path}")
    try:
        # Validate resume collection
        if resume_collection is None:
            logger.error("ChromaDB resume collection not initialized.")
            return {"status": "error", "message": "ChromaDB resume collection not initialized."}
        # Parse resume (reuse parse_resume tool)
        parse_result = parse_resume(project_name, file_path)
        if parse_result.get("status") != "success":
            logger.error(f"parse_resume failed: {parse_result.get('message')}")
            return {"status": "error", "message": f"parse_resume failed: {parse_result.get('message')}"}
        raw_text = parse_result.get("raw_text", "")
        metadata = parse_result.get("extracted_fields", {})
        # Chunking
        text = raw_text.strip()
        if not text:
            logger.error("No text to chunk for indexing.")
            return {"status": "error", "message": "No text to chunk for indexing."}
        # Section/heading-based chunking
        import re
        headings = [
            r"SUMMARY", r"PROFILE", r"OBJECTIVE", r"EXPERIENCE", r"WORK HISTORY", r"EMPLOYMENT",
            r"EDUCATION", r"SKILLS", r"PROJECTS", r"CERTIFICATION", r"CERTIFICATIONS", r"ACHIEVEMENTS",
            r"PUBLICATIONS", r"AWARDS", r"LANGUAGES", r"INTERESTS", r"PERSONAL", r"REFERENCES"
        ]
        # Regex: heading must be at line start, all caps, possibly with spaces, at least 5 chars
        heading_regex = re.compile(rf"^({'|'.join(headings)})[\s:]*$", re.MULTILINE | re.IGNORECASE)
        matches = list(heading_regex.finditer(text))
        chunks = []
        if matches:
            logger.info(f"Found {len(matches)} section headings; using section-based chunking.")
            positions = [m.start() for m in matches] + [len(text)]
            for i in range(len(positions)-1):
                chunk = text[positions[i]:positions[i+1]].strip()
                if chunk:
                    chunks.append(chunk)
            logger.info(f"Section-based chunking produced {len(chunks)} chunks.")
        else:
            # Fallback to fixed-size chunking
            logger.info("No section headings found; using fixed-size chunking.")
            start = 0
            while start < len(text):
                end = min(start + chunk_size, len(text))
                chunk = text[start:end]
                if chunk.strip():
                    chunks.append(chunk)
                start += chunk_size - chunk_overlap
            logger.info(f"Fixed-size chunking produced {len(chunks)} chunks.")
        # Embedding (Ollama, nomic-embed-text)
        try:
            import requests
            ollama_url = "http://localhost:11434/api/embeddings"
            embeddings = []
            errors = []
            for i, chunk in enumerate(chunks):
                payload = {"model": "nomic-embed-text", "prompt": chunk}
                try:
                    resp = requests.post(ollama_url, json=payload, timeout=30)
                    resp.raise_for_status()
                    emb = resp.json().get("embedding")
                    if not emb:
                        raise ValueError("No embedding in Ollama response.")
                    embeddings.append(emb)
                except Exception as emb_e:
                    logger.error(f"Embedding failed for chunk {i}: {emb_e}")
                    errors.append(f"Chunk {i}: {emb_e}")
                    embeddings.append([0.0]) # Placeholder for failed embedding
            # Store in ChromaDB
            ids = [str(uuid.uuid4()) for _ in chunks]
            import json as _json
            def _make_scalar(val):
                if isinstance(val, (str, int, float, bool)) or val is None:
                    return val
                try:
                    return _json.dumps(val, ensure_ascii=False)
                except Exception:
                    return str(val)
            metadatas = [{
                "file_path": file_path,
                "project_name": project_name,
                "chunk_index": i,
                **{k: _make_scalar(v) for k, v in metadata.items()}
            } for i in range(len(chunks))]
            try:
                resume_collection.add(
                    documents=chunks,
                    embeddings=embeddings,
                    metadatas=metadatas,
                    ids=ids
                )
            except Exception as chroma_e:
                logger.error(f"ChromaDB add failed: {chroma_e}\nMetadata sample: {metadatas[0] if metadatas else 'N/A'}")
                return {"status": "error", "message": f"ChromaDB add failed: {chroma_e}"}
            logger.info(f"Indexed {len(chunks)} chunks for resume: {file_path}")
            return {
                "status": "success",
                "message": f"Indexed {len(chunks)} chunks for resume.",
                "chunks_indexed": len(chunks),
                "embedding_errors": errors,
                "file_path": file_path
            }
        except Exception as embed_e:
            logger.error(f"Embedding or ChromaDB storage failed: {embed_e}")
            return {"status": "error", "message": f"Embedding or storage failed: {embed_e}"}
    except Exception as main_e:
        logger.error(f"Critical error in index_resume: {main_e}", exc_info=True)
        return {"status": "error", "message": f"Critical error: {main_e}"}
# --- End Resume Indexing Tool ---

# --- Resume Semantic Search Tool ---
# Recruiting tool moved to recruiting_mcp_server – decorator removed
def semantic_search_in_resumes(query: str, max_results: int = 5) -> dict:
    """
    Performs semantic search over indexed resume chunks using vector similarity (Ollama nomic-embed-text).
    Args:
        query: The natural language search query (e.g., job description, keywords).
        max_results: Maximum number of top matches to return (default 5).
    Returns:
        Dict with a 'results' list sorted by similarity, each item containing:
            - 'file_path': Path to the original resume file
            - 'chunk_content': The matching chunk of resume text
            - 'similarity_score': Higher is more similar
            - 'metadata': All extracted fields and pyresparser data
    """
    logger.info(f"Starting semantic_search_in_resumes for query: {query}")
    try:
        if resume_collection is None:
            logger.error("ChromaDB resume collection not initialized.")
            return {"results": [], "error": "ChromaDB resume collection not initialized."}
        # Embed the query
        try:
            import requests
            ollama_url = "http://localhost:11434/api/embeddings"
            payload = {"model": "nomic-embed-text", "prompt": query}
            resp = requests.post(ollama_url, json=payload, timeout=30)
            resp.raise_for_status()
            query_emb = resp.json().get("embedding")
            if not query_emb:
                raise ValueError("No embedding in Ollama response.")
        except Exception as emb_e:
            logger.error(f"Failed to embed query: {emb_e}")
            return {"results": [], "error": f"Embedding failed: {emb_e}"}
        # Query ChromaDB
        try:
            search = resume_collection.query(
                query_embeddings=[query_emb],
                n_results=max_results,
                include=["documents", "embeddings", "metadatas", "distances"]
            )
            results = []
            docs = search.get("documents", [[]])[0]
            metadatas = search.get("metadatas", [[]])[0]
            distances = search.get("distances", [[]])[0]
            for i in range(len(docs)):
                similarity_score = 1.0 / (1.0 + distances[i]) if distances[i] is not None else 0.0
                metadata = metadatas[i] if i < len(metadatas) else {}
                results.append({
                    "file_path": metadata.get("file_path", "Unknown"),
                    "chunk_content": docs[i],
                    "similarity_score": round(similarity_score, 4),
                    "metadata": metadata
                })
            logger.info(f"Semantic search returned {len(results)} results for query: {query}")
            return {"results": results}
        except Exception as search_e:
            logger.error(f"ChromaDB search failed: {search_e}")
            return {"results": [], "error": f"ChromaDB search failed: {search_e}"}
    except Exception as main_e:
        logger.error(f"Critical error in semantic_search_in_resumes: {main_e}", exc_info=True)
        return {"results": [], "error": f"Critical error: {main_e}"}
# --- End Resume Semantic Search Tool ---

# --- Hybrid Search Recruiting Tool ---
# Recruiting tool moved to recruiting_mcp_server – decorator removed
def hybrid_search_recruiting_data(
    query: str,
    core_technologies: List[str],
    max_results: int = 10,
    tech_aliases: Optional[Dict[str, List[str]]] = None,
    required_technologies: Optional[List[str]] = None,
    score_weights: Optional[Dict[str, float]] = None,
    fuzzy_threshold: int = 85
) -> dict:
    """
    Hybrid search for recruiting: combines semantic search, keyword/alias, and optional fuzzy matching to rank candidate resumes for a job description.

    Args:
        query: Job description or search text.
        core_technologies: List of required core skills/technologies.
        max_results: Maximum number of candidates to return (default 10).
        tech_aliases: Dictionary mapping each core tech to a list of aliases/synonyms.
        required_technologies: Subset of core_technologies that must be present (optional).
        score_weights: Dict with weights for 'semantic' and 'tech' scores (default: {'semantic': 0.6, 'tech': 0.4}).
        fuzzy_threshold: Fuzzy match threshold (0-100, default 85). Set >100 to disable fuzzy.
    Returns:
        Dict with ranked candidates, each with file_path, final_score, semantic_score, tech_match_score, matched/missing techs, matched_aliases, and chunk_content.
    """
    logger.info(f"[hybrid_search_recruiting_data] Starting hybrid search for query: {query}")
    try:
        # Validate thefuzz
        if fuzzy_threshold <= 100 and fuzz is None:
            logger.error("thefuzz library not installed but fuzzy matching requested.")
            return {"status": "error", "message": "thefuzz not installed. Run: pip install thefuzz[spd]"}
        # Validate args
        if not query or not core_technologies:
            return {"status": "error", "message": "Must provide both query and core_technologies."}
        tech_aliases = tech_aliases or {}
        required_technologies = set(required_technologies) if required_technologies else set()
        score_weights = score_weights or {"semantic": 0.6, "tech": 0.4}
        # 1. Semantic search
        semantic_results = semantic_search_in_resumes(query, max_results=max_results * 2)
        if not semantic_results or "results" not in semantic_results:
            logger.error("No semantic search results returned.")
            return {"status": "error", "message": "Semantic search failed or returned no results."}
        candidates = []
        for cand in semantic_results["results"]:
            chunk_text = cand.get("chunk_content", "")
            norm_text = chunk_text.lower()
            matched_techs = set()
            missing_techs = set(core_technologies)
            matched_aliases = {}
            for tech in core_technologies:
                aliases = set([tech.lower()])
                if tech_aliases.get(tech):
                    aliases.update([a.lower() for a in tech_aliases[tech]])
                found = False
                # Alias/keyword match (whole word/phrase)
                for alias in aliases:
                    if alias in norm_text:
                        matched_techs.add(tech)
                        matched_aliases[tech] = alias
                        found = True
                        break
                # Fuzzy match if not found and enabled
                if not found and fuzzy_threshold <= 100 and fuzz is not None:
                    for alias in aliases:
                        # Token set ratio for phrase/word matching
                        score = fuzz.token_set_ratio(alias, norm_text)
                        if score >= fuzzy_threshold:
                            matched_techs.add(tech)
                            matched_aliases[tech] = f"fuzzy:{alias}"
                            found = True
                            break
                if found:
                    missing_techs.discard(tech)
            tech_match_score = len(matched_techs) / len(core_technologies) if core_technologies else 0.0
            # Filter by required_technologies
            if required_technologies and not required_technologies.issubset(matched_techs):
                continue
            candidates.append({
                "file_path": cand.get("file_path", "Unknown"),
                "chunk_content": chunk_text,
                "semantic_score": float(cand.get("similarity_score", 0.0)) * 100,
                "tech_match_score": tech_match_score * 100,
                "core_tech_matched": sorted(list(matched_techs)),
                "core_tech_missing": sorted(list(missing_techs)),
                "matched_aliases_found": matched_aliases
            })
        # 4. Score combination/ranking
        for cand in candidates:
            cand["final_score"] = (
                score_weights.get("semantic", 0.6) * cand["semantic_score"] +
                score_weights.get("tech", 0.4) * cand["tech_match_score"]
            )
        candidates.sort(key=lambda c: c["final_score"], reverse=True)
        logger.info(f"[hybrid_search_recruiting_data] Returning {min(len(candidates), max_results)} candidates.")
        return {
            "status": "success",
            "results": candidates[:max_results],
            "total_candidates": len(candidates)
        }
    except Exception as e:
        logger.error(f"[hybrid_search_recruiting_data] Critical error: {e}", exc_info=True)
        return {"status": "error", "message": f"Critical error: {e}"}
# --- End Hybrid Search Recruiting Tool ---

# --- Job Description Parsing Tool ---
# Recruiting tool moved to recruiting_mcp_server – decorator removed
def parse_job_description(file_path: str = None, raw_text: str = None) -> dict:
    """
    Parses a job description (JD) file or raw text and extracts structured requirements.
    Args:
        file_path: Path to the job description file (PDF, DOCX, TXT, or MD). Optional if raw_text is provided.
        raw_text: Raw job description text (if not using a file).
    Returns:
        Dict with status, extracted fields (title, required_skills, min_experience, education, certifications, keywords),
        the raw JD text, and error/message if any.
    """
    logger.info(f"Starting parse_job_description for file: {file_path} raw_text: {bool(raw_text)}")
    try:
        if not file_path and not raw_text:
            return {"status": "error", "message": "Must provide either file_path or raw_text."}
        # Load text from file if needed
        text = raw_text or ""
        if file_path:
            try:
                ext = str(file_path).lower().split('.')[-1]
                if ext == "pdf":
                    from pdfminer.high_level import extract_text
                    text = extract_text(file_path)
                elif ext in ("docx", "doc"):
                    import docx
                    doc = docx.Document(file_path)
                    text = "\n".join([p.text for p in doc.paragraphs])
                elif ext in ("txt", "md"):
                    with open(file_path, "r", encoding="utf-8") as f:
                        text = f.read()
                else:
                    return {"status": "error", "message": f"Unsupported file extension: {ext}"}
            except Exception as file_e:
                logger.error(f"Failed to read job description file: {file_e}")
                return {"status": "error", "message": f"Failed to read file: {file_e}"}
        if not text or not text.strip():
            return {"status": "error", "message": "No text found in job description."}
        # Extraction logic
        import re
        import spacy
        nlp = spacy.load("en_core_web_sm")
        doc = nlp(text)
        # Title: first non-empty line, or from heading
        lines = [l.strip() for l in text.splitlines() if l.strip()]
        title = lines[0] if lines else "Unknown"
        # Skills: look for keywords (simple list, could be improved with a skills DB)
        SKILL_KEYWORDS = [
            "python", "java", "c++", "c#", "javascript", "typescript", "sql", "aws", "azure", "docker",
            "kubernetes", "react", "node", "django", "flask", "fastapi", "tensorflow", "pytorch", "nlp",
            "machine learning", "deep learning", "data science", "rest api", "microservices", "git",
            "linux", "unix", "agile", "scrum", "jira", "html", "css", "sass", "spark", "hadoop",
            "gcp", "cloud", "devops", "testing", "selenium", "ci/cd", "graphql", "mongodb", "postgresql"
        ]
        found_skills = set()
        for skill in SKILL_KEYWORDS:
            if re.search(rf"\b{re.escape(skill)}\b", text, re.IGNORECASE):
                found_skills.add(skill)
        # Experience: look for patterns like 'X+ years', 'at least X years', etc.
        exp_matches = re.findall(r"(\d+)[+ ]+years?", text, re.IGNORECASE)
        min_experience = int(exp_matches[0]) if exp_matches else None
        # Education: look for degree keywords
        EDUCATION_KEYWORDS = [
            "bachelor", "master", "phd", "b.sc", "m.sc", "b.e", "m.e", "b.tech", "m.tech", "mba", "degree", "diploma"
        ]
        found_education = []
        for kw in EDUCATION_KEYWORDS:
            if re.search(rf"\b{re.escape(kw)}\b", text, re.IGNORECASE):
                found_education.append(kw)
        # Certifications: look for keywords
        CERT_KEYWORDS = ["aws certified", "azure certified", "pmp", "scrum master", "gcp certified", "oracle certified", "cfa", "cissp", "security+", "network+"]
        found_certs = []
        for cert in CERT_KEYWORDS:
            if re.search(rf"\b{re.escape(cert)}\b", text, re.IGNORECASE):
                found_certs.append(cert)
        # Keywords: extract top 10 nouns/adjectives by frequency
        from collections import Counter
        tokens = [t.lemma_.lower() for t in doc if t.pos_ in ("NOUN", "ADJ") and not t.is_stop and len(t.text) > 2]
        keyword_counts = Counter(tokens)
        keywords = [kw for kw, _ in keyword_counts.most_common(10)]
        logger.info(f"parse_job_description extracted: title={title}, skills={found_skills}, min_exp={min_experience}, education={found_education}, certs={found_certs}, keywords={keywords}")
        return {
            "status": "success",
            "title": title,
            "required_skills": sorted(found_skills),
            "min_experience": min_experience,
            "education": found_education,
            "certifications": found_certs,
            "keywords": keywords,
            "raw_text": text
        }
    except Exception as main_e:
        logger.error(f"Critical error in parse_job_description: {main_e}", exc_info=True)
        return {"status": "error", "message": f"Critical error: {main_e}"}
# --- End Job Description Parsing Tool ---

# --- MCP Server Startup Banner and Request Logging ---
logger.info("""
==============================\nMCP Server (v0.2.0) STARTUP\n==============================\nProject: Local Project Context Server\nPython: %s\nPlatform: %s\nWorking Directory: %s\nPID: %s\n------------------------------\nIf you are running this via an IDE (VS Code Copilot Agent, Windsurf Cascade, etc.),\nthis banner should appear in your IDE's MCP server logs.\n------------------------------\n""" % (platform.python_version(), platform.platform(), os.getcwd(), os.getpid()))

# --- Log all incoming MCP requests (for stdio debugging) ---
try:
    if hasattr(mcp, "_handle"):
        # Patch request logging if possible (example: mcp._handle = logging_handle)
        logger.info("MCP request logging is ENABLED (all incoming requests will be logged).")
    else:
        logger.info("FastMCP _handle attribute not present; skipping request logging patch.")
except Exception as e:
    logger.warning(f"Could not patch FastMCP for request logging: {e}")

# === Core Tools ===

@mcp.tool()
def get_server_info() -> dict:
    """
    Returns basic information about the server environment.
    """
    logger.info("DEBUG: Entered get_server_info function.") # <-- ADD THIS
    logger.info(f"DEBUG: Inside get_server_info - Type of mcp: {type(mcp)}, Has version attr: {hasattr(mcp, 'version')}") # <-- ADD THIS
    try: # <-- ADD TRY/EXCEPT
        logger.info("Executing get_server_info tool")
        info = {
            "message": "MCP Server is running via stdio!",
            "server_version": "0.2.0", # Accessing the version here
            "python_version": platform.python_version(),
            "system": platform.system(),
            "release": platform.release()
        }
        logger.info(f"Returning server info: {info}")
        return info
    except AttributeError as e: # <-- Catch the specific error
        logger.error(f"DEBUG: Caught AttributeError inside get_server_info: {e}", exc_info=True)
        # Optionally, return an error dictionary instead of letting it crash the tool execution
        return {"status": "error", "message": f"AttributeError during tool execution: {e}"}
    except Exception as e: # <-- Catch any other unexpected error
         logger.error(f"DEBUG: Caught unexpected Exception inside get_server_info: {e}", exc_info=True)
         return {"status": "error", "message": f"Unexpected error during tool execution: {e}"}

# === File System Interaction Tools ===

@mcp.tool()
def list_project_markdown_files(project_name: str) -> list[str]:
    """
    Lists all Markdown (.md) files found recursively within the documentation
    directory of the specified project ('ParentBuddy' or 'DreamApp').
    Returns a list of absolute file paths. Includes extra debug logging.
    """
    logger.info(f"Executing list_project_markdown_files for project: {project_name}")
    if project_name not in PROJECT_DOC_PATHS:
        logger.error(f"Invalid project name '{project_name}'. Valid names are: {list(PROJECT_DOC_PATHS.keys())}")
        return []
    doc_path = PROJECT_DOC_PATHS[project_name]
    markdown_files = []
    try:
        abs_path_str = str(doc_path.resolve())
        logger.debug(f"Checking configured path object: {doc_path}")
        logger.debug(f"Absolute path string resolved to: {abs_path_str}")
        if doc_path.exists():
            if doc_path.is_dir():
                logger.debug(f"Attempting to run rglob('*.md') on directory: {abs_path_str}")
                files_found_count = 0
                try:
                    for file_path_obj in doc_path.rglob('*.md'):
                        if file_path_obj.is_file():
                            resolved_file_path = str(file_path_obj.resolve())
                            markdown_files.append(resolved_file_path)
                            files_found_count += 1
                    logger.debug(f"rglob('*.md') iteration finished. Found {files_found_count} files.")
                except PermissionError as rglob_pe:
                     logger.error(f"Permission denied during rglob file search within {abs_path_str}", exc_info=True)
                     return []
                except Exception as rglob_e:
                    logger.error(f"Error during rglob file search within {abs_path_str}: {rglob_e}", exc_info=True)
                    return []
            else:
                logger.error(f"Path exists but is not a directory: {abs_path_str}")
                return []
        else:
            logger.error(f"Configured path does not exist: {abs_path_str}")
            return []
        logger.info(f"Found {len(markdown_files)} markdown files for project '{project_name}'.")
        return markdown_files
    except PermissionError as pe:
        logger.error(f"Permission denied while initially accessing path: {doc_path}", exc_info=True)
        return []
    except Exception as e:
        logger.error(f"An unexpected error occurred listing files for '{project_name}': {e}", exc_info=True)
        return []

@mcp.tool()
def read_project_file(absolute_file_path: str) -> dict:
    """
    Reads and returns the full text content of the specified project
    documentation file using its absolute path. Validates path against allowed roots.
    Use paths returned by list_project_markdown_files.

    Args:
        absolute_file_path: The absolute path to the documentation file to read.

    Returns:
        A dictionary with:
            - 'status': 'success' or 'error'.
            - 'file_path': The file path that was attempted.
            - 'content': The file content as a string (empty if error).
            - 'message': Informational or error message.
    """
    logger.info(f"Executing read_project_file for path: {absolute_file_path}")
    try:
        file_path = pathlib.Path(absolute_file_path)
        is_safe_path = False
        for root_path in PROJECT_DOC_PATHS.values():
            try:
                if file_path.resolve().is_relative_to(root_path.resolve()):
                    is_safe_path = True
                    break
            except ValueError:
                continue
            except Exception as path_e:
                logger.warning(f"Path comparison error for {file_path} against {root_path}: {path_e}")
                continue
        if not is_safe_path:
            logger.error(f"Access denied: Path '{absolute_file_path}' is outside configured documentation directories.")
            return {
                "status": "error",
                "file_path": absolute_file_path,
                "content": "",
                "message": "Access denied: file is outside allowed documentation directories."
            }
        if not file_path.is_file():
            logger.error(f"Path exists but is not a file: {absolute_file_path}")
            return {
                "status": "error",
                "file_path": absolute_file_path,
                "content": "",
                "message": "Path exists but is not a file."
            }
        content = file_path.read_text(encoding='utf-8')
        logger.info(f"Successfully read {len(content)} characters from: {absolute_file_path}")
        return {
            "status": "success",
            "file_path": absolute_file_path,
            "content": content,
            "message": f"Read {len(content)} characters."
        }
    except FileNotFoundError:
        logger.error(f"File not found during read attempt: {absolute_file_path}")
        return {
            "status": "error",
            "file_path": absolute_file_path,
            "content": "",
            "message": "File not found."
        }
    except PermissionError:
        logger.error(f"Permission denied while reading file: {absolute_file_path}", exc_info=True)
        return {
            "status": "error",
            "file_path": absolute_file_path,
            "content": "",
            "message": "Permission denied."
        }
    except Exception as e:
        logger.error(f"An unexpected error occurred reading file '{absolute_file_path}': {e}", exc_info=True)
        return {
            "status": "error",
            "file_path": absolute_file_path,
            "content": "",
            "message": f"Unexpected error: {e}"
        }


# === Source File Reading Tool ===

@mcp.tool()
def read_project_source_file(absolute_file_path: str) -> dict:
    """
    Reads and returns the full text content of the specified Dart source file using its absolute path.
    Validates the path against allowed project source roots (see PROJECT_SRC_PATHS).
    Provides robust logging and error handling. Only allows access to files within configured source directories.

    Args:
        absolute_file_path: The absolute path to the Dart source file to read.

    Returns:
        A dictionary with:
            - 'status': 'success' or 'error'.
            - 'file_path': The file path that was attempted.
            - 'content': The file content as a string (empty if error).
            - 'message': Informational or error message.
    """
    logger.info(f"Executing read_project_source_file for path: {absolute_file_path}")
    try:
        file_path = pathlib.Path(absolute_file_path)
        is_safe_path = False
        for root_path in PROJECT_SRC_PATHS.values():
            try:
                if file_path.resolve().is_relative_to(root_path.resolve()):
                    is_safe_path = True
                    break
            except Exception as e:
                logger.debug(f"Path resolution error during source path check: {e}")
        if not is_safe_path:
            logger.warning(f"Access denied for source file outside allowed roots: {absolute_file_path}")
            return {
                "status": "error",
                "file_path": absolute_file_path,
                "content": "",
                "message": "Access denied: file is outside allowed project source directories."
            }
        if not file_path.exists():
            logger.error(f"Source file not found: {absolute_file_path}")
            return {
                "status": "error",
                "file_path": absolute_file_path,
                "content": "",
                "message": "File not found."
            }
        if not file_path.is_file():
            logger.error(f"Path is not a file: {absolute_file_path}")
            return {
                "status": "error",
                "file_path": absolute_file_path,
                "content": "",
                "message": "Path is not a file."
            }
        try:
            content = file_path.read_text(encoding='utf-8')
            logger.info(f"Successfully read {len(content)} characters from Dart source: {absolute_file_path}")
            return {
                "status": "success",
                "file_path": absolute_file_path,
                "content": content,
                "message": f"Read {len(content)} characters."
            }
        except Exception as file_read_e:
            logger.error(f"Error reading Dart source file: {absolute_file_path}: {file_read_e}", exc_info=True)
            return {
                "status": "error",
                "file_path": absolute_file_path,
                "content": "",
                "message": f"Error reading file: {file_read_e}"
            }
    except Exception as e:
        logger.error(f"Unexpected error in read_project_source_file for '{absolute_file_path}': {e}", exc_info=True)
        return {
            "status": "error",
            "file_path": absolute_file_path,
            "content": "",
            "message": f"Unexpected error: {e}"
        }

# === Dart Source Code Analysis Tool ===

@mcp.tool()
def get_dart_definition(project_name: str, file_path: str, symbol_name: str) -> dict:
    """
    Finds and extracts the source code definition for a specific Dart class
    or function within a given file using tree-sitter parsing.

    Args:
        project_name: The name of the project ('ParentBuddy' or 'DreamApp').
                      Used primarily for logging/context here, path validation
                      relies on read_project_source_file.
        file_path: The absolute path to the .dart source file.
        symbol_name: The exact name of the class or function to find.

    Returns:
        A dictionary containing:
        - 'status': 'success', 'error', or 'not_found'.
        - 'symbol_name': The requested symbol name.
        - 'file_path': The source file path.
        - 'definition': The extracted source code block (string) or None.
        - 'start_line': The starting line number (1-based) of the definition, or None.
        - 'end_line': The ending line number (1-based) of the definition, or None.
        - 'message': An informational or error message.
    """
    logger.info(f"Executing get_dart_definition for symbol '{symbol_name}' in file: {file_path} (Project: {project_name})")

    # Check if tree-sitter setup was successful
    if not DART_LANGUAGE or not dart_parser:
        error_msg = "Tree-sitter for Dart is not available. Cannot parse code."
        logger.error(error_msg)
        return {"status": "error", "symbol_name": symbol_name, "file_path": file_path, "definition": None, "start_line": None, "end_line": None, "message": error_msg}

    # Validate inputs
    if not project_name or not file_path or not symbol_name:
         return {"status": "error", "symbol_name": symbol_name, "file_path": file_path, "definition": None, "start_line": None, "end_line": None, "message": "Project name, file path, and symbol name are required."}

    # Read the source file content using the existing tool for path validation
    file_result = read_project_source_file(file_path)
    if file_result["status"] != "success" or not file_result["content"]:
        # Error message already logged by read_project_source_file
        return {"status": "error", "symbol_name": symbol_name, "file_path": file_path, "definition": None, "start_line": None, "end_line": None, "message": f"Failed to read or access source file: {file_path}"}

    file_content = file_result["content"]

    try:
        # Parse the source code using tree-sitter
        tree = dart_parser.parse(bytes(file_content, "utf8"))
        root_node = tree.root_node

        # --- Tree-sitter Query to find class or function definition ---
        # This query looks for class definitions or function definitions (including methods)
        # where the identifier matches the desired symbol_name.
        # It captures the entire definition node (@definition) and the name node (@name).
        query_string = f"""
        (class_definition
          name: (identifier) @name
          (#eq? @name "{symbol_name}")
        ) @definition

        (function_signature
          name: (identifier) @name
          (#eq? @name "{symbol_name}")
          body: _
        ) @definition_sig
         (function_body) @definition_body ;; Find function body separately if needed

        (method_signature ;; Also find methods within classes
          name: (identifier) @name
          (#eq? @name "{symbol_name}")
          body: _
        ) @definition ;; Capture method signature + body typically

        ;; Add other potential top-level definitions if needed (e.g., variables)
        """

        query = DART_LANGUAGE.query(query_string)
        captures = query.captures(root_node)

        # Find the first match for the definition
        definition_node = None
        for node, capture_name in captures:
            if capture_name == 'definition': # Prioritize capturing the whole definition block
                definition_node = node
                break
            # Fallback or alternative logic might be needed depending on query results

        if definition_node:
            start_byte = definition_node.start_byte
            end_byte = definition_node.end_byte
            definition_text = file_content[start_byte:end_byte]

            # Get line numbers (0-based from tree-sitter, convert to 1-based)
            start_line = definition_node.start_point[0] + 1
            end_line = definition_node.end_point[0] + 1

            logger.info(f"Found definition for '{symbol_name}' in {file_path} (Lines {start_line}-{end_line})")
            return {
                "status": "success",
                "symbol_name": symbol_name,
                "file_path": file_path,
                "definition": definition_text,
                "start_line": start_line,
                "end_line": end_line,
                "message": "Definition found successfully."
            }
        else:
            logger.warning(f"Definition for symbol '{symbol_name}' not found in {file_path}")
            return {
                "status": "not_found",
                "symbol_name": symbol_name,
                "file_path": file_path,
                "definition": None,
                "start_line": None,
                "end_line": None,
                "message": f"Symbol '{symbol_name}' definition not found in the file."
            }

    except Exception as parse_e:
        logger.error(f"Error parsing file {file_path} or querying for symbol '{symbol_name}': {parse_e}", exc_info=True)
        return {"status": "error", "symbol_name": symbol_name, "file_path": file_path, "definition": None, "start_line": None, "end_line": None, "message": f"Error during parsing or query: {parse_e}"}

# === Resume Parsing Tool ===

import re
import mimetypes
import chardet

try:
    import docx
except ImportError:
    docx = None
try:
    from pdfminer.high_level import extract_text as extract_pdf_text
except ImportError:
    extract_pdf_text = None
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

@mcp.tool()
def parse_resume(project_name: str, file_path: str) -> dict:
    """
    Parses a resume file (PDF, DOCX, TXT, or MD) and extracts raw text plus key fields using pyresparser (spaCy) if available.

    Args:
        project_name: The name of the recruiting project (must be in PROJECT_RESUME_PATHS).
        file_path: The absolute path to the resume file.

    Returns:
        A dictionary containing:
        - 'status': 'success', 'error', or 'unsupported'.
        - 'file_path': The file path attempted.
        - 'raw_text': The extracted text (or empty string on error).
        - 'extracted_fields':
            - 'emails': List of email addresses found.
            - 'phone_numbers': List of phone numbers found.
            - 'potential_skills': List of matched skills (from a default list).
            - 'names': List of candidate name guesses (best effort).
            - 'summary': First 2-3 lines/sentences as a preview.
            - 'pyresparser': Dict of extractions from pyresparser (if available):
                - 'name', 'email', 'mobile_number', 'skills', 'education', 'total_experience', 'designation', etc.
        - 'message': Informational or error message.
    """
    logger.info(f"Executing parse_resume for project: {project_name}, file: {file_path}")
    # --- Validate project and file path ---
    if project_name not in PROJECT_RESUME_PATHS:
        logger.error(f"Invalid project name '{project_name}'. Valid: {list(PROJECT_RESUME_PATHS.keys())}")
        return {"status": "error", "file_path": file_path, "raw_text": "", "extracted_fields": {}, "message": f"Invalid project name: {project_name}"}
    resume_root = PROJECT_RESUME_PATHS[project_name]
    try:
        abs_resume_root = resume_root.resolve()
        abs_file_path = pathlib.Path(file_path).resolve()
        if not abs_file_path.is_file() or not str(abs_file_path).startswith(str(abs_resume_root)):
            logger.error(f"File path {abs_file_path} is not a file or not within allowed resume root {abs_resume_root}")
            return {"status": "error", "file_path": file_path, "raw_text": "", "extracted_fields": {}, "message": "File not found or not allowed."}
    except Exception as e:
        logger.error(f"Error validating file path: {e}")
        return {"status": "error", "file_path": file_path, "raw_text": "", "extracted_fields": {}, "message": f"Path validation error: {e}"}

    # --- Detect file type ---
    ext = abs_file_path.suffix.lower()
    mimetype, _ = mimetypes.guess_type(str(abs_file_path))
    raw_text = ""
    error = None
    try:
        if ext == ".pdf" or (mimetype and "pdf" in mimetype):
            if extract_pdf_text:
                logger.info("Using pdfminer.six to extract PDF text.")
                try:
                    raw_text = extract_pdf_text(str(abs_file_path))
                except Exception as pdfminer_e:
                    logger.warning(f"pdfminer failed: {pdfminer_e}. Trying PyPDF2.")
                    if PyPDF2:
                        try:
                            with open(abs_file_path, "rb") as f:
                                reader = PyPDF2.PdfReader(f)
                                raw_text = "\n".join(page.extract_text() or '' for page in reader.pages)
                        except Exception as pypdf_e:
                            error = f"PDF extraction failed with both pdfminer and PyPDF2: {pypdf_e}"
                    else:
                        error = f"pdfminer failed and PyPDF2 not available: {pdfminer_e}"
            elif PyPDF2:
                logger.info("Using PyPDF2 to extract PDF text.")
                try:
                    with open(abs_file_path, "rb") as f:
                        reader = PyPDF2.PdfReader(f)
                        raw_text = "\n".join(page.extract_text() or '' for page in reader.pages)
                except Exception as pypdf_e:
                    error = f"PDF extraction failed with PyPDF2: {pypdf_e}"
            else:
                error = "No PDF text extraction library available. Please install pdfminer.six or PyPDF2."
        elif ext == ".docx" or (mimetype and "word" in mimetype):
            if docx:
                logger.info("Using python-docx to extract DOCX text.")
                try:
                    doc = docx.Document(str(abs_file_path))
                    raw_text = "\n".join([para.text for para in doc.paragraphs])
                except Exception as docx_e:
                    error = f"DOCX extraction failed: {docx_e}"
            else:
                error = "python-docx not available. Please install it."
        elif ext in [".txt", ".md"] or (mimetype and ("text" in mimetype or "markdown" in mimetype)):
            logger.info("Reading as plain text/markdown.")
            try:
                with open(abs_file_path, "rb") as f:
                    raw_bytes = f.read()
                    detected = chardet.detect(raw_bytes)
                    encoding = detected.get("encoding") or "utf-8"
                    raw_text = raw_bytes.decode(encoding, errors="replace")
            except Exception as txt_e:
                error = f"Text/Markdown extraction failed: {txt_e}"
        else:
            logger.error(f"Unsupported file type: {ext}")
            return {"status": "unsupported", "file_path": file_path, "raw_text": "", "extracted_fields": {}, "message": f"Unsupported file type: {ext}"}
    except Exception as extract_e:
        logger.error(f"General extraction error: {extract_e}")
        return {"status": "error", "file_path": file_path, "raw_text": "", "extracted_fields": {}, "message": f"Extraction error: {extract_e}"}

    if error:
        logger.error(error)
        return {"status": "error", "file_path": file_path, "raw_text": raw_text, "extracted_fields": {}, "message": error}
    if not raw_text.strip():
        logger.warning("No text extracted from file.")
        return {"status": "error", "file_path": file_path, "raw_text": "", "extracted_fields": {}, "message": "No text extracted from file."}

    # --- Extract fields ---
    extracted_fields = {}
    # Try pyresparser for advanced extraction
    pyresparser_data = None
    try:
        from pyresparser import ResumeParser
        import spacy
        # Check if spaCy model is downloaded
        try:
            nlp = spacy.load('en_core_web_sm')
        except OSError:
            logger.warning("spaCy model 'en_core_web_sm' not found. Please run: python -m spacy download en_core_web_sm")
            nlp = None
        if nlp:
            pyresparser_data = ResumeParser(str(abs_file_path)).get_extracted_data()
            logger.info(f"pyresparser extracted fields: {list(pyresparser_data.keys())}")
    except ImportError as e:
        logger.warning(f"pyresparser or spaCy not installed: {e}")
    except Exception as e:
        logger.warning(f"pyresparser failed: {e}")
    if pyresparser_data:
        extracted_fields['pyresparser'] = pyresparser_data
    # Emails
    email_pattern = r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+"
    extracted_fields["emails"] = re.findall(email_pattern, raw_text)
    # Phone numbers (basic international and US patterns)
    phone_pattern = r"(?:\+\d{1,3}[\s-]?)?(?:\(?\d{3}\)?[\s-]?)?\d{3}[\s-]?\d{4}"
    extracted_fields["phone_numbers"] = re.findall(phone_pattern, raw_text)
    # Potential skills (default set, easily extendable)
    default_skills = [
        "python", "java", "c++", "javascript", "typescript", "react", "node", "sql", "aws",
        "docker", "kubernetes", "linux", "git", "html", "css", "flutter", "dart", "go", "ruby",
        "php", "azure", "gcp", "tensorflow", "pytorch", "spark", "hadoop", "scala", "project management",
        "agile", "scrum", "jira", "ci/cd", "devops", "testing", "selenium", "graphql", "mongodb", "postgresql"
    ]
    found_skills = set()
    for skill in default_skills:
        # Simple word-boundary match, case-insensitive
        if re.search(rf"\\b{re.escape(skill)}\\b", raw_text, re.IGNORECASE):
            found_skills.add(skill)
    extracted_fields["potential_skills"] = list(found_skills)
    # Names (best effort: from file name and first lines)
    file_stem = abs_file_path.stem
    first_lines = raw_text.splitlines()[:5]
    name_guesses = []
    # Try to find lines that look like "Firstname Lastname" (very basic)
    name_regex = r"^[A-Z][a-z]+\s+[A-Z][a-z]+$"
    for line in first_lines:
        if re.match(name_regex, line.strip()):
            name_guesses.append(line.strip())
    # Add file stem if it looks like a name
    if re.match(name_regex, file_stem.replace('_', ' ').replace('-', ' ')):
        name_guesses.append(file_stem.replace('_', ' ').replace('-', ' '))
    extracted_fields["names"] = name_guesses
    # Summary: first 2-3 non-empty lines/sentences
    summary_lines = [l.strip() for l in raw_text.splitlines() if l.strip()][:3]
    extracted_fields["summary"] = " ".join(summary_lines)

    logger.info(f"parse_resume succeeded for {file_path}")
    return {
        "status": "success",
        "file_path": file_path,
        "raw_text": raw_text,
        "extracted_fields": extracted_fields,
        "message": f"Resume parsed successfully. Extracted {len(raw_text)} characters."
    }

# === Database Schema Tool ===

@mcp.tool()
def get_drift_schema(project_name: str) -> dict:
    """
    Parses .drift files within a specified project to extract the database schema
    (tables, columns, types, and table-level constraints) defined using Drift syntax.
    Note: Uses regex parsing, may have limitations with complex syntax.

    Args:
        project_name: The name of the project ('ParentBuddy' or 'DreamApp')
                      to search for .drift files within.

    Returns:
        A dictionary containing:
        - 'status': 'success', 'error', or 'not_found'.
        - 'project_name': The requested project name.
        - 'schema': A dictionary where keys are table names and values are dicts with:
            - 'columns': list of column dicts (name, type, constraints)
            - 'table_constraints': list of table-level constraints (as strings)
        - 'files_parsed': A list of absolute paths to the .drift files that were parsed.
        - 'message': An informational or error message.
    """
    logger.info(f"Executing get_drift_schema for project: {project_name}")

    if project_name not in PROJECT_DRIFT_PATHS:
        logger.error(f"Invalid project name '{project_name}' for Drift schema extraction. Valid names are: {list(PROJECT_DRIFT_PATHS.keys())}")
        return {"status": "error", "project_name": project_name, "schema": {}, "files_parsed": [], "message": f"Invalid project name: {project_name}"}

    search_root_path = PROJECT_DRIFT_PATHS[project_name]
    drift_files_found = []
    schema_extracted = {}  # {table_name: {columns: [...], table_constraints: [...]}}

    try:
        abs_search_path_str = str(search_root_path.resolve())
        logger.debug(f"Searching for .drift files starting from: {abs_search_path_str}")

        if not search_root_path.is_dir():
            logger.error(f"Drift search path is not a valid directory: {abs_search_path_str}")
            return {"status": "error", "project_name": project_name, "schema": {}, "files_parsed": [], "message": f"Search path not a directory: {abs_search_path_str}"}

        files_found_count = 0
        for file_path_obj in search_root_path.rglob('*.drift'):
            if file_path_obj.is_file():
                resolved_file_path = str(file_path_obj.resolve())
                drift_files_found.append(resolved_file_path)
                files_found_count += 1
        logger.info(f"Found {files_found_count} .drift files for project '{project_name}'.")

        if not drift_files_found:
            return {"status": "not_found", "project_name": project_name, "schema": {}, "files_parsed": [], "message": "No .drift files found in the project."}

    except PermissionError as pe:
        logger.error(f"Permission denied while searching for .drift files in {abs_search_path_str}", exc_info=True)
        return {"status": "error", "project_name": project_name, "schema": {}, "files_parsed": [], "message": f"Permission error during search: {pe}"}
    except Exception as search_e:
        logger.error(f"Error searching for .drift files in {abs_search_path_str}: {search_e}", exc_info=True)
        return {"status": "error", "project_name": project_name, "schema": {}, "files_parsed": [], "message": f"Error during search: {search_e}"}

    # Regex patterns
    table_pattern = re.compile(r"(\w+)\s*:\s*TABLE\s+WITH\s+Columns\s*\((.*?)\)\s*END", re.IGNORECASE | re.DOTALL | re.MULTILINE)
    column_pattern = re.compile(r"(\w+)\s+([\w<>]+(?:\(\d+\))?)\s*(.*)", re.IGNORECASE)
    # Table-level constraints: CHECK, UNIQUE, FOREIGN KEY, PRIMARY KEY (multi-column)
    constraint_pattern = re.compile(r"^(CHECK|UNIQUE|FOREIGN\s+KEY|PRIMARY\s+KEY)\s*\(.*", re.IGNORECASE)

    files_parsed_successfully = []
    for file_path in drift_files_found:
        logger.debug(f"Parsing drift file: {file_path}")
        try:
            content = pathlib.Path(file_path).read_text(encoding='utf-8')

            for table_match in table_pattern.finditer(content):
                table_name = table_match.group(1).strip()
                columns_block = table_match.group(2).strip()
                logger.debug(f"Found table definition for: {table_name}")

                if table_name in schema_extracted:
                    logger.warning(f"Table '{table_name}' defined multiple times (found again in {file_path}). Overwriting previous definition.")
                schema_extracted[table_name] = {"columns": [], "table_constraints": []}

                lines = columns_block.split('\n')
                for line in lines:
                    line_strip = line.strip()
                    if not line_strip or line_strip.startswith('--'):
                        continue
                    # Table-level constraint?
                    constraint_match = constraint_pattern.match(line_strip)
                    if constraint_match:
                        schema_extracted[table_name]["table_constraints"].append(line_strip)
                        logger.debug(f"  - Found table constraint: {line_strip}")
                        continue
                    col_match = column_pattern.match(line_strip)
                    if col_match:
                        col_name = col_match.group(1).strip()
                        col_type = col_match.group(2).strip().upper()
                        col_constraints = col_match.group(3).strip() if col_match.group(3) else None
                        schema_extracted[table_name]["columns"].append({
                            "name": col_name,
                            "type": col_type,
                            "constraints": col_constraints if col_constraints else None
                        })
                        logger.debug(f"  - Found column: {col_name} {col_type} {col_constraints}")
                    else:
                        logger.warning(f"Could not parse line in table '{table_name}' as column or constraint: {line_strip}")

            files_parsed_successfully.append(file_path)

        except FileNotFoundError:
            logger.error(f"Drift file not found during read attempt (should not happen after rglob): {file_path}")
        except PermissionError:
            logger.error(f"Permission denied while reading drift file: {file_path}", exc_info=True)
        except Exception as parse_e:
            logger.error(f"An unexpected error occurred parsing drift file '{file_path}': {parse_e}", exc_info=True)

    if not schema_extracted:
        message = "Successfully parsed drift files, but no table definitions were found."
        logger.warning(message)
        return {"status": "not_found", "project_name": project_name, "schema": {}, "files_parsed": files_parsed_successfully, "message": message}
    else:
        message = f"Successfully extracted schema for {len(schema_extracted)} tables from {len(files_parsed_successfully)} drift files."
        logger.info(message)
        return {"status": "success", "project_name": project_name, "schema": schema_extracted, "files_parsed": files_parsed_successfully, "message": message}

    return {"status": "error", "project_name": project_name, "schema": {}, "files_parsed": [], "message": "Unknown error in get_drift_schema."}

@mcp.tool()
def get_sqlite_schema(project_name: str, file_path: str) -> dict:
    """
    Extracts database schema information from a Drift (.dart) file defining database tables
    for a Flutter application. Parses the file to find table definitions, columns, types, 
    and constraints.
    
    Args:
        project_name: The name of the project ('ParentBuddy' or 'DreamApp').
                      Used primarily for logging/context, path validation
                      relies on read_project_source_file.
        file_path: The absolute path to the Drift (.dart) file containing database schema definitions.
    
    Returns:
        A dictionary containing:
        - 'status': 'success', 'error', or 'empty'.
        - 'file_path': The source file path.
        - 'tables': A list of table definitions, each containing:
            - 'name': Table name
            - 'columns': List of columns, each with 'name', 'type', and 'constraints'
        - 'message': An informational or error message.
    """
    logger.info(f"Executing get_sqlite_schema for file: {file_path} (Project: {project_name})")
    
    # Validate inputs
    if not project_name or not file_path:
        error_msg = "Project name and file path are required."
        logger.error(error_msg)
        return {"status": "error", "file_path": file_path, "tables": [], "message": error_msg}
    
    # Read the source file content using the existing tool for path validation
    file_result = read_project_source_file(file_path)
    if file_result["status"] != "success" or not file_result["content"]:
        # Error message already logged by read_project_source_file
        return {"status": "error", "file_path": file_path, "tables": [], "message": f"Failed to read or access source file: {file_path}"}

    file_content = file_result["content"]
    
    # Check if this looks like a Drift file (contains class with @DataClassName or extends Table)
    if 'extends Table' not in file_content and '@DataClassName' not in file_content and 'drift:' not in file_content:
        warn_msg = f"File {file_path} does not appear to be a Drift database schema file."
        logger.warning(warn_msg)
        return {"status": "error", "file_path": file_path, "tables": [], "message": warn_msg}
    
    try:
        # Parse the file to extract table definitions
        tables = []
        table_class_pattern = r'class\s+([A-Za-z0-9_]+)\s+extends\s+Table\s*{([^}]+)}'  
        table_matches = re.finditer(table_class_pattern, file_content, re.DOTALL)
        
        for table_match in table_matches:
            table_name = table_match.group(1)
            table_body = table_match.group(2)
            
            # Extract columns
            columns = []
            column_pattern = r'TextColumn\s+get\s+([A-Za-z0-9_]+)\s*=>\s*text\(\)(.+?);|' \
                          r'IntColumn\s+get\s+([A-Za-z0-9_]+)\s*=>\s*integer\(\)(.+?);|' \
                          r'RealColumn\s+get\s+([A-Za-z0-9_]+)\s*=>\s*real\(\)(.+?);|' \
                          r'BoolColumn\s+get\s+([A-Za-z0-9_]+)\s*=>\s*boolean\(\)(.+?);|' \
                          r'DateTimeColumn\s+get\s+([A-Za-z0-9_]+)\s*=>\s*dateTime\(\)(.+?);|' \
                          r'BlobColumn\s+get\s+([A-Za-z0-9_]+)\s*=>\s*blob\(\)(.+?);'
            
            column_matches = re.finditer(column_pattern, table_body, re.DOTALL)
            
            for column_match in column_matches:
                # Find which column type was matched
                col_name = None
                col_type = None
                col_constraints = []
                
                for i, group in enumerate(column_match.groups()):
                    if i % 2 == 0 and group:  # Name groups are at even indices
                        col_name = group
                        col_type = ['text', 'integer', 'real', 'boolean', 'dateTime', 'blob'][i//2]
                    elif i % 2 == 1 and group:  # Constraint groups are at odd indices
                        # Extract constraints
                        if '.required()' in group:
                            col_constraints.append('NOT NULL')
                        if '.unique()' in group:
                            col_constraints.append('UNIQUE')
                        if '.primaryKey()' in group or '.autoIncrement()' in group:
                            col_constraints.append('PRIMARY KEY')
                        if '.references(' in group:
                            # Try to extract foreign key reference
                            ref_pattern = r'.references\(([^)]+)\)'
                            ref_match = re.search(ref_pattern, group)
                            if ref_match:
                                col_constraints.append(f"REFERENCES {ref_match.group(1)}")
                
                if col_name and col_type:
                    columns.append({
                        'name': col_name,
                        'type': col_type,
                        'constraints': col_constraints
                    })
            
            tables.append({
                'name': table_name,
                'columns': columns
            })
        
        if not tables:
            # Try alternative pattern for drift: notation
            drift_table_pattern = r'drift:\s*\"CREATE\s+TABLE\s+([A-Za-z0-9_]+)\s*\(([^\)]+)\)\"'
            drift_matches = re.finditer(drift_table_pattern, file_content, re.DOTALL | re.IGNORECASE)
            
            for drift_match in drift_matches:
                table_name = drift_match.group(1)
                columns_text = drift_match.group(2)
                
                columns = []
                column_pattern = r'\s*([A-Za-z0-9_]+)\s+([A-Za-z0-9_]+)\s*(.*)'
                column_matches = re.finditer(column_pattern, columns_text)
                
                for column_match in column_matches:
                    if len(column_match.groups()) >= 2:
                        col_name = column_match.group(1)
                        col_type = column_match.group(2)
                        col_constraints = []
                        
                        if len(column_match.groups()) >= 3 and column_match.group(3):
                            constraints_text = column_match.group(3).strip()
                            if constraints_text:
                                # Simple constraint extraction
                                if 'NOT NULL' in constraints_text.upper():
                                    col_constraints.append('NOT NULL')
                                if 'UNIQUE' in constraints_text.upper():
                                    col_constraints.append('UNIQUE')
                                if 'PRIMARY KEY' in constraints_text.upper():
                                    col_constraints.append('PRIMARY KEY')
                                if 'REFERENCES' in constraints_text.upper():
                                    ref_pattern = r'REFERENCES\s+([A-Za-z0-9_\.]+)'
                                    ref_match = re.search(ref_pattern, constraints_text, re.IGNORECASE)
                                    if ref_match:
                                        col_constraints.append(f"REFERENCES {ref_match.group(1)}")
                        
                        columns.append({
                            'name': col_name,
                            'type': col_type,
                            'constraints': col_constraints
                        })
                
                tables.append({
                    'name': table_name,
                    'columns': columns
                })
        
        if not tables:
            return {
                "status": "empty",
                "file_path": file_path,
                "tables": [],
                "message": "No database schema definitions found in the file."
            }
        
        logger.info(f"Successfully extracted schema for {len(tables)} tables from {file_path}")
        return {
            "status": "success",
            "file_path": file_path,
            "tables": tables,
            "message": f"Successfully extracted schema for {len(tables)} tables."
        }
        
    except Exception as e:
        logger.error(f"Error extracting database schema from {file_path}: {e}", exc_info=True)
        return {
            "status": "error",
            "file_path": file_path,
            "tables": [],
            "message": f"Error extracting schema: {e}"
        }

# === Documentation Parsing Tool ===

@mcp.tool()
def parse_markdown_to_tokens(markdown_text: str) -> list[dict]:
    """
    Parses the provided Markdown text into a list of token objects using
    markdown-it-py. Each token represents a structural element.
    Returns a list of tokens serialized as dictionaries.
    """
    logger.info(f"Executing parse_markdown_to_tokens for text length: {len(markdown_text)}")
    if not isinstance(markdown_text, str):
        logger.error("Invalid input: markdown_text must be a string.")
        return []
    try:
        tokens: list[Token] = md_parser.parse(markdown_text)
        serialized_tokens = [token.as_dict() for token in tokens] # Simpler serialization
        logger.info(f"Successfully parsed markdown into {len(serialized_tokens)} tokens.")
        return serialized_tokens
    except Exception as e:
        logger.error(f"An unexpected error occurred during markdown parsing: {e}", exc_info=True)
        return []

# === Vector DB Indexing and Search Tools ===

@mcp.tool()
def index_project_documentation(project_name: str) -> dict:
    """
    Indexes documentation for a project. Reads files, chunks, generates embeddings
    via local Ollama (nomic-embed-text), stores in ChromaDB. Returns status.
    REQUIRES OLLAMA TO BE RUNNING.
    """
    logger.info(f"Starting indexing for project: {project_name} using Ollama/nomic-embed-text")
    if not chroma_client or not collection:
         return {"status": "error", "message": "ChromaDB client/collection not initialized."}
    if project_name not in PROJECT_DOC_PATHS:
        return {"status": "error", "message": f"Invalid project name '{project_name}'"}

    files_processed, chunks_added, errors_occurred = 0, 0, 0
    embedding_model_name = "nomic-embed-text"
    try:
        markdown_files = list_project_markdown_files(project_name)
        if not markdown_files:
             return {"status": "success", "message": f"No markdown files found for '{project_name}'.", "files_processed": 0, "chunks_added": 0, "errors": 0}
        logger.info(f"Found {len(markdown_files)} files to index for '{project_name}'.")

        for file_path in markdown_files:
            logger.debug(f"Processing file: {file_path}")
            file_had_error = False
            try:
                content = read_project_file(file_path)
                if not content: continue
                text_chunks = [chunk for chunk in content.split('\n\n') if chunk.strip()]
                if not text_chunks: continue
                logger.debug(f"Split content into {len(text_chunks)} chunks for {file_path}")

                chunk_embeddings, chunk_ids, chunk_documents, chunk_metadatas = [], [], [], []
                for i, chunk_text in enumerate(text_chunks):
                    chunk_id = f"{file_path}_{i}"
                    try:
                        embedding_response = ollama.embeddings(model=embedding_model_name, prompt=chunk_text)
                        chunk_embeddings.append(embedding_response['embedding'])
                        chunk_ids.append(chunk_id)
                        chunk_documents.append(chunk_text)
                        chunk_metadatas.append({"source": file_path, "project": project_name})
                    except Exception as embed_error:
                        logger.error(f"Ollama embedding failed for chunk {i} of {file_path}: {embed_error}", exc_info=True)
                        errors_occurred += 1; file_had_error = True; continue

                if chunk_ids:
                    try:
                        collection.add(ids=chunk_ids, embeddings=chunk_embeddings, documents=chunk_documents, metadatas=chunk_metadatas)
                        chunks_added += len(chunk_ids)
                        logger.debug(f"Added {len(chunk_ids)} chunks from {file_path} to ChromaDB.")
                    except Exception as chroma_error:
                         logger.error(f"ChromaDB add failed for {len(chunk_ids)} chunks from {file_path}: {chroma_error}", exc_info=True)
                         errors_occurred += len(chunk_ids); file_had_error = True
            except Exception as file_proc_error:
                logger.error(f"Failed to process file {file_path}: {file_proc_error}", exc_info=True)
                errors_occurred += 1; file_had_error = True; continue
            if not file_had_error: files_processed += 1
    except Exception as main_error:
        logger.error(f"Critical error during indexing for '{project_name}': {main_error}", exc_info=True)
        return {"status": "error", "message": f"Critical error: {main_error}", "files_processed": files_processed, "chunks_added": chunks_added, "errors": errors_occurred}

    status_msg = f"Indexing complete for '{project_name}'. Files OK: {files_processed}. Chunks added: {chunks_added}. Errors: {errors_occurred}."
    logger.info(status_msg)
    return {"status": "success", "message": status_msg, "files_processed": files_processed, "chunks_added": chunks_added, "errors": errors_occurred}


@mcp.tool(name="keyword_search_in_documentation")
async def keyword_search_in_documentation(query: str, project_name: str | None = None, max_results: int = 10) -> dict:
    """
    Performs keyword search within indexed documentation chunks using case-insensitive matching.

    Args:
        query: The keyword(s) to search for.
        project_name: Optional. Filter results to 'ParentBuddy' or 'DreamApp'.
        max_results: Optional. Max number of results (default 10).

    Returns:
        Dict with 'results' list, each item having 'file_path' and 'chunk_content'.
    """
    logger.info(f"Keyword search: query='{query}', project='{project_name}', max_results={max_results}")
    if not collection: return {"results": []} # Guard clause

    search_results = []
    try:
        where_filter = {}
        if project_name and project_name in PROJECT_DOC_PATHS:
            where_filter["project"] = project_name
        elif project_name:
             logger.warning(f"Invalid project_name '{project_name}' ignored in keyword search.")

        candidate_limit = max(max_results * 5, 20)
        logger.debug(f"ChromaDB candidate query: where={where_filter}, limit={candidate_limit}")
        results = collection.get(
            where=where_filter if where_filter else None,
            limit=candidate_limit,
            include=["metadatas", "documents"]
        )

        retrieved_count = len(results.get('ids', []))
        logger.info(f"ChromaDB returned {retrieved_count} candidates for keyword search.")
        if retrieved_count > 0:
            query_lower = query.lower()
            for i in range(retrieved_count):
                doc_content = results['documents'][i]
                if query_lower in doc_content.lower():
                    metadata = results['metadatas'][i]
                    search_results.append({
                        "file_path": metadata.get("source", "Unknown Path"),
                        "chunk_content": doc_content
                    })
                    if len(search_results) >= max_results: break
        logger.info(f"Formatted {len(search_results)} keyword results for: '{query}'")
    except Exception as e:
        logger.error(f"Error during keyword search for '{query}': {e}", exc_info=True)
        return {"results": []}
    return {"results": search_results}


@mcp.tool(name="semantic_search_in_documentation")
async def semantic_search_in_documentation(query: str, project_name: str | None = None, max_results: int = 5) -> dict:
    """
    Performs semantic search using vector similarity (Ollama nomic-embed-text).

    Args:
        query: The natural language query.
        project_name: Optional. Filter results to 'ParentBuddy' or 'DreamApp'.
        max_results: Optional. Max number of results (default 5).

    Returns:
        Dict with 'results' list sorted by similarity, each item having
        'file_path', 'chunk_content', and 'similarity_score'.
    """
    logger.info(f"Semantic search: query='{query}', project='{project_name}', max_results={max_results}")
    if not collection: return {"results": []} # Guard clause

    search_results = []
    embedding_model_name = "nomic-embed-text"
    try:
        # 1. Embed Query
        logger.debug(f"Generating Ollama embedding for query: '{query}'")
        try:
            embedding_response = ollama.embeddings(model=embedding_model_name, prompt=query)
            query_embedding = embedding_response['embedding']
            logger.debug("Query embedding generated.")
        except Exception as embed_error:
            logger.error(f"Ollama embedding failed for query '{query}': {embed_error}", exc_info=True)
            return {"results": []}

        # 2. Build Filter
        where_filter = {}
        if project_name and project_name in PROJECT_DOC_PATHS:
            where_filter["project"] = project_name
        elif project_name:
             logger.warning(f"Invalid project_name '{project_name}' ignored in semantic search.")

        # 3. Query ChromaDB
        logger.debug(f"Querying ChromaDB '{COLLECTION_NAME}'. Filter: {where_filter}, N_results: {max_results}")
        results = collection.query(
            query_embeddings=[query_embedding],
            n_results=max_results,
            where=where_filter if where_filter else None,
            include=["metadatas", "documents", "distances"]
        )

        # 4. Process Results
        retrieved_count = len(results.get('ids', [[]])[0])
        logger.info(f"ChromaDB vector query returned {retrieved_count} results.")
        if retrieved_count > 0:
            ids, distances, metadatas, documents = (results.get(k, [[]])[0] for k in ['ids', 'distances', 'metadatas', 'documents'])
            for i in range(retrieved_count):
                similarity_score = 1.0 / (1.0 + distances[i]) # Simple inverse distance
                search_results.append({
                    "file_path": metadatas[i].get("source", "Unknown Path"),
                    "chunk_content": documents[i],
                    "similarity_score": round(similarity_score, 4)
                })
        logger.info(f"Formatted {len(search_results)} semantic results for: '{query}'")
    except Exception as e:
        logger.error(f"Error during semantic search for '{query}': {e}", exc_info=True)
        return {"results": []}
    return {"results": search_results}

# === NEW Granular Documentation Parsing Tools ===

@mcp.tool()
def debug_markdown_tokens(project_name: str, file_path: str) -> dict:
    """
    Reads a markdown file from a project's documentation directory and returns all tokens parsed by markdown-it-py.
    Useful for debugging markdown parsing issues and understanding the token structure.

    Args:
        project_name: The name of the project ('ParentBuddy', 'DreamApp', etc.).
        file_path: The absolute path to the markdown (.md) file.

    Returns:
        A dictionary with:
            - 'status': 'success' or 'error'.
            - 'file_path': The file path that was parsed.
            - 'tokens': A list of token dictionaries (or empty list on error).
            - 'message': Informational or error message.
    """
    logger.info(f"Executing debug_markdown_tokens for project '{project_name}', file '{file_path}'")
    try:
        # Validate project
        if project_name not in PROJECT_DOC_PATHS:
            msg = f"Invalid project name: {project_name}"
            logger.error(msg)
            return {"status": "error", "file_path": file_path, "tokens": [], "message": msg}
        # Validate file path is within allowed doc root
        doc_root = PROJECT_DOC_PATHS[project_name]
        abs_file_path = pathlib.Path(file_path).resolve()
        if not abs_file_path.is_file() or not str(abs_file_path).startswith(str(doc_root.resolve())):
            msg = f"File path {abs_file_path} is not a file or not within allowed doc root {doc_root}"
            logger.error(msg)
            return {"status": "error", "file_path": file_path, "tokens": [], "message": msg}
        # Read file
        try:
            content = abs_file_path.read_text(encoding="utf-8")
        except Exception as read_e:
            msg = f"Failed to read file: {read_e}"
            logger.error(msg)
            return {"status": "error", "file_path": file_path, "tokens": [], "message": msg}
        # Parse tokens
        try:
            tokens: list[Token] = md_parser.parse(content)
            serialized_tokens = [token.as_dict() for token in tokens]
            logger.info(f"Parsed {len(serialized_tokens)} tokens from markdown file: {file_path}")
            return {"status": "success", "file_path": file_path, "tokens": serialized_tokens, "message": "Tokens parsed successfully."}
        except Exception as parse_e:
            msg = f"Markdown parsing failed: {parse_e}"
            logger.error(msg)
            return {"status": "error", "file_path": file_path, "tokens": [], "message": msg}
    except Exception as main_e:
        msg = f"Unexpected error in debug_markdown_tokens: {main_e}"
        logger.error(msg, exc_info=True)
        return {"status": "error", "file_path": file_path, "tokens": [], "message": msg}

@mcp.tool()
def extract_user_flow(project_name: str, feature_name: str) -> dict:
    """
    Extracts the step-by-step user flow for a specified feature
    from the Markdown documentation of a given project. Searches
    for headings related to the feature and extracts subsequent ordered lists.

    Args:
        project_name: The name of the project ('ParentBuddy' or 'DreamApp').
        feature_name: A descriptive name of the feature whose flow is needed
                      (e.g., 'Login', 'Add Dream', 'Create Reminder').

    Returns:
        A dictionary containing:
        - 'status': 'success' or 'error'.
        - 'feature_name': The requested feature name.
        - 'source_file': The file path where the flow was found (or None).
        - 'flow_steps': A list of strings, each representing a step in the flow (or []).
        - 'message': An informational or error message.
    """
    logger.info(f"Executing extract_user_flow for project '{project_name}', feature '{feature_name}'")
    if not project_name or not feature_name:
        return {"status": "error", "feature_name": feature_name, "source_file": None, "flow_steps": [], "message": "Project name and feature name required."}
    if project_name not in PROJECT_DOC_PATHS:
        return {"status": "error", "feature_name": feature_name, "source_file": None, "flow_steps": [], "message": f"Invalid project name: {project_name}"}

    try:
        markdown_files = list_project_markdown_files(project_name)
        if not markdown_files:
            return {"status": "success", "feature_name": feature_name, "source_file": None, "flow_steps": [], "message": "No documentation files found."}

        result = {"status": "success", "feature_name": feature_name, "source_file": None, "flow_steps": [], "message": "User flow not found."}
        # NOTE: Regex patterns are examples and likely need refinement for specific doc structures
        heading_pattern = re.compile(r"^\s*#{2,}\s+.*(?:Flow|Scenario).*" + re.escape(feature_name) + r".*$", re.IGNORECASE | re.MULTILINE)
        list_item_pattern = re.compile(r"^\s*(?:\d+\.|[-*+])\s+(.*)$")

        for file_path in markdown_files:
            logger.debug(f"Searching for flow in file: {file_path}")
            try:
                content = read_project_file(file_path) # Relies on read_project_file for safety
                if not content: continue
                match = heading_pattern.search(content)
                if match:
                    logger.info(f"Found potential feature heading for '{feature_name}' in {file_path}")
                    result["source_file"] = file_path
                    content_after_heading = content[match.end():]
                    next_heading_match = re.search(r"^\s*#{1,}", content_after_heading, re.MULTILINE)
                    section_content = content_after_heading[:next_heading_match.start()] if next_heading_match else content_after_heading
                    flow_steps = []
                    in_list = False
                    for line in section_content.strip().split('\n'):
                        list_match = list_item_pattern.match(line.strip())
                        if list_match:
                            in_list = True
                            flow_steps.append(list_match.group(1).strip())
                        elif in_list: break # Assume list ended
                    if flow_steps:
                        logger.info(f"Extracted {len(flow_steps)} steps for '{feature_name}' from {file_path}")
                        result["flow_steps"] = flow_steps
                        result["message"] = "User flow extracted successfully."
                        return result # Found it, return immediately
                    else:
                         logger.warning(f"Found heading for '{feature_name}' in {file_path}, but no subsequent list items detected.")
            except Exception as file_read_e:
                logger.error(f"Error reading or processing file {file_path} for flow extraction: {file_read_e}", exc_info=True)
        # If loop completes without finding the flow
        logger.warning(f"Could not find user flow documentation for feature: '{feature_name}' in project '{project_name}'.")
        return result
    except Exception as e:
        logger.error(f"Unexpected error during user flow extraction for '{feature_name}': {e}", exc_info=True)
        return {"status": "error", "feature_name": feature_name, "source_file": None, "flow_steps": [], "message": f"Unexpected error: {e}"}

@mcp.tool()
def get_wireframe_details(project_name: str, screen_or_component_name: str) -> dict:
    """
    Retrieves descriptive details about a specific screen or UI component
    from Markdown wireframe/design documentation for the specified project.
    Searches for relevant headings or sections.

    Args:
        project_name: The name of the project ('ParentBuddy' or 'DreamApp').
        screen_or_component_name: The name of the screen or component to find details for
                                   (e.g., 'Dream Journal Screen', 'Login Button').

    Returns:
        A dictionary containing:
        - 'status': 'success' or 'error'.
        - 'screen_or_component': The requested screen/component name.
        - 'source_file': The file path where the details were found (or None).
        - 'details': A string containing the extracted description/details (or None).
        - 'message': An informational or error message.
    """
    logger.info(f"Executing get_wireframe_details for project '{project_name}', target '{screen_or_component_name}'")
    if not project_name or not screen_or_component_name:
        return {"status": "error", "screen_or_component": screen_or_component_name, "source_file": None, "details": None, "message": "Project name and screen/component name required."}
    if project_name not in PROJECT_DOC_PATHS:
        return {"status": "error", "screen_or_component": screen_or_component_name, "source_file": None, "details": None, "message": f"Invalid project name: {project_name}"}

    try:
        markdown_files = list_project_markdown_files(project_name)
        if not markdown_files:
            return {"status": "success", "screen_or_component": screen_or_component_name, "source_file": None, "details": None, "message": "No documentation files found."}

        result = {"status": "success", "screen_or_component": screen_or_component_name, "source_file": None, "details": None, "message": "Details not found."}
        # NOTE: Regex pattern is an example, adapt based on actual doc format
        heading_pattern = re.compile(r"^\s*#{1,}\s+.*(?:Screen|Component|Wireframe)?:?\s*" + re.escape(screen_or_component_name) + r".*$", re.IGNORECASE | re.MULTILINE)

        for file_path in markdown_files:
            logger.debug(f"Searching for wireframe details in file: {file_path}")
            try:
                content = read_project_file(file_path) # Relies on read_project_file for safety
                if not content: continue
                match = heading_pattern.search(content)
                if match:
                    logger.info(f"Found potential section for '{screen_or_component_name}' in {file_path}")
                    result["source_file"] = file_path
                    content_after_heading = content[match.end():]
                    next_heading_match = re.search(r"^\s*#{1,}", content_after_heading, re.MULTILINE)
                    section_content = content_after_heading[:next_heading_match.start()].strip() if next_heading_match else content_after_heading.strip()
                    if section_content:
                        logger.info(f"Extracted details for '{screen_or_component_name}' from {file_path}")
                        result["details"] = section_content
                        result["message"] = "Details extracted successfully."
                        return result # Found it, return immediately
                    else:
                        logger.warning(f"Found heading for '{screen_or_component_name}' in {file_path}, but no content followed.")
            except Exception as file_read_e:
                logger.error(f"Error reading or processing file {file_path} for wireframe details: {file_read_e}", exc_info=True)
        # If loop completes without finding details
        logger.warning(f"Could not find wireframe/design details for: '{screen_or_component_name}' in project '{project_name}'.")
        return result
    except Exception as e:
        logger.error(f"Unexpected error during wireframe detail retrieval for '{screen_or_component_name}': {e}", exc_info=True)
        return {"status": "error", "screen_or_component": screen_or_component_name, "source_file": None, "details": None, "message": f"Unexpected error: {e}"}

@mcp.tool()
def get_style_guide_spec(project_name: str, style_property: str) -> dict:
    """
    Extracts a specific style guide property value (e.g., 'Primary Color')
    from the Markdown documentation for the specified project.
    Searches for key-value pairs or definitions.

    Args:
        project_name: The name of the project ('ParentBuddy' or 'DreamApp').
        style_property: The name of the style property to retrieve
                        (e.g., 'Primary Color', 'Body Font').

    Returns:
        A dictionary containing:
        - 'status': 'success' or 'error'.
        - 'property_name': The requested style property name.
        - 'source_file': The file path where the value was found (or None).
        - 'value': The extracted value of the style property (or None).
        - 'message': An informational or error message.
    """
    logger.info(f"Executing get_style_guide_spec for project '{project_name}', property '{style_property}'")
    if not project_name or not style_property:
        return {"status": "error", "property_name": style_property, "source_file": None, "value": None, "message": "Project name and style property required."}
    if project_name not in PROJECT_DOC_PATHS:
        return {"status": "error", "property_name": style_property, "source_file": None, "value": None, "message": f"Invalid project name: {project_name}"}

    try:
        markdown_files = list_project_markdown_files(project_name)
        if not markdown_files:
            return {"status": "success", "property_name": style_property, "source_file": None, "value": None, "message": "No documentation files found."}

        result = {"status": "success", "property_name": style_property, "source_file": None, "value": None, "message": "Style property value not found."}
        # NOTE: Regex pattern is an example, adapt based on actual doc format (e.g., `**Prop:** Value`, `- Prop: Value`)
        kv_pattern = re.compile(r"^\s*(?:\*\*|[-*+])?\s*" + re.escape(style_property) + r"\s*[:\-]\s*(.+)$", re.IGNORECASE | re.MULTILINE)

        for file_path in markdown_files:
             # Optional: Prioritize files like 'StyleGuide.md'
            is_priority_file = any(term in file_path.lower() for term in ["style", "design", "theme"])
            logger.debug(f"Searching for style property '{style_property}' in file: {file_path}" + (" (priority)" if is_priority_file else ""))
            try:
                content = read_project_file(file_path) # Relies on read_project_file for safety
                if not content: continue
                match = kv_pattern.search(content)
                if match:
                    extracted_value = match.group(1).strip()
                    # Optional: Add further cleanup for extracted value if needed
                    logger.info(f"Found style property '{style_property}' with value '{extracted_value}' in {file_path}")
                    result["source_file"] = file_path
                    result["value"] = extracted_value
                    result["message"] = "Style property value extracted successfully."
                    return result # Found it, return immediately
                # Add logic here to parse tables or other structures if needed
            except Exception as file_read_e:
                logger.error(f"Error reading or processing file {file_path} for style specs: {file_read_e}", exc_info=True)
        # If loop completes without finding the property
        logger.warning(f"Could not find style guide specification for property: '{style_property}' in project '{project_name}'.")
        return result
    except Exception as e:
        logger.error(f"Unexpected error during style guide spec retrieval for '{style_property}': {e}", exc_info=True)
        return {"status": "error", "property_name": style_property, "source_file": None, "value": None, "message": f"Unexpected error: {e}"}

# === Explicitly Start the MCP Stdio Loop ===
if __name__ == "__main__":
    import argparse, json
    parser = argparse.ArgumentParser(description="MCP Server CLI mode")
    parser.add_argument('--cli-tool', type=str, help='Tool name to run (e.g., hybrid_search_recruiting_data)')
    parser.add_argument('--cli-args-file', type=str, help='Path to a JSON file containing arguments for the tool')
    args, unknown = parser.parse_known_args()
    if args.cli_tool and args.cli_args_file:
        # CLI mode: run a single tool and exit
        try:
            tool_func = getattr(sys.modules[__name__], args.cli_tool, None)
            if not tool_func:
                print(json.dumps({"status": "error", "message": f"Tool '{args.cli_tool}' not found."}))
                sys.exit(2)
            try:
                with open(args.cli_args_file, 'r', encoding='utf-8') as f:
                    tool_args = json.load(f)
            except FileNotFoundError:
                print(json.dumps({"status": "error", "message": f"Argument file not found: {args.cli_args_file}"}))
                sys.exit(4)
            except json.JSONDecodeError as json_e:
                print(json.dumps({"status": "error", "message": f"Invalid JSON in argument file {args.cli_args_file}: {json_e}"}))
                sys.exit(5)

            result = tool_func(**tool_args)
            print(json.dumps(result))
            sys.exit(0)
        except Exception as e:
            print(json.dumps({"status": "error", "message": str(e)}))
            sys.exit(3)
    # Default: stdio loop
    logger.info("Script execution finished defining tools. Attempting to start MCP stdio loop...")
    try:
        mcp.run_stdio()
        logger.info("FastMCP stdio loop finished.")
    except AttributeError:
        logger.warning("mcp.run_stdio() not found. Trying mcp.run() as fallback...")
        try:
            mcp.run()
            logger.info("FastMCP run() loop finished.")
        except Exception as run_err:
            logger.error(f"Failed to execute mcp.run(): {run_err}", exc_info=True)
            sys.exit(1)
    except Exception as stdio_err:
        logger.error(f"Error during MCP stdio loop execution: {stdio_err}", exc_info=True)
        sys.exit(1)